"""
ExonEditor v7
Tela inicial com seleção de modo:
  • Automático — API NCBI (busca direta no GenBank)
  • Manual     — Upload .docx + texto copiado do GenBank
Ambos os modos com: barra visual, análise de códons, geração de documentos.
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, colorchooser
import re, os, itertools, threading
import urllib.request, urllib.parse, urllib.error
import json, time
from docx import Document
from docx.shared import RGBColor, Pt


# ══════════════════════════════════════════════════════════════
# CÓDIGO GENÉTICO
# ══════════════════════════════════════════════════════════════

CODIGO_GENETICO = {
    'TTT':'Phe','TTC':'Phe','TTA':'Leu','TTG':'Leu',
    'CTT':'Leu','CTC':'Leu','CTA':'Leu','CTG':'Leu',
    'ATT':'Ile','ATC':'Ile','ATA':'Ile','ATG':'Met',
    'GTT':'Val','GTC':'Val','GTA':'Val','GTG':'Val',
    'TCT':'Ser','TCC':'Ser','TCA':'Ser','TCG':'Ser',
    'CCT':'Pro','CCC':'Pro','CCA':'Pro','CCG':'Pro',
    'ACT':'Thr','ACC':'Thr','ACA':'Thr','ACG':'Thr',
    'GCT':'Ala','GCC':'Ala','GCA':'Ala','GCG':'Ala',
    'TAT':'Tyr','TAC':'Tyr','TAA':'Stop','TAG':'Stop',
    'CAT':'His','CAC':'His','CAA':'Gln','CAG':'Gln',
    'AAT':'Asn','AAC':'Asn','AAA':'Lys','AAG':'Lys',
    'GAT':'Asp','GAC':'Asp','GAA':'Glu','GAG':'Glu',
    'TGT':'Cys','TGC':'Cys','TGA':'Stop','TGG':'Trp',
    'CGT':'Arg','CGC':'Arg','CGA':'Arg','CGG':'Arg',
    'AGT':'Ser','AGC':'Ser','AGA':'Arg','AGG':'Arg',
    'GGT':'Gly','GGC':'Gly','GGA':'Gly','GGG':'Gly',
}

AA_SINONIMOS = {
    'A':'Ala','R':'Arg','N':'Asn','D':'Asp','C':'Cys',
    'Q':'Gln','E':'Glu','G':'Gly','H':'His','I':'Ile',
    'L':'Leu','K':'Lys','M':'Met','F':'Phe','P':'Pro',
    'S':'Ser','T':'Thr','W':'Trp','Y':'Tyr','V':'Val',
    'ALA':'Ala','ARG':'Arg','ASN':'Asn','ASP':'Asp','CYS':'Cys',
    'GLN':'Gln','GLU':'Glu','GLY':'Gly','HIS':'His','ILE':'Ile',
    'LEU':'Leu','LYS':'Lys','MET':'Met','PHE':'Phe','PRO':'Pro',
    'SER':'Ser','THR':'Thr','TRP':'Trp','TYR':'Tyr','VAL':'Val',
}

def normalizar_aa(e):
    return AA_SINONIMOS.get(e.strip().upper())

def codons_para_aa(aa3):
    return [c for c, a in CODIGO_GENETICO.items() if a == aa3]

def construir_cds(seq, exons):
    exons_ord = sorted(exons, key=lambda x: x['inicio'])
    cds, mapa = '', []
    for ex in exons_ord:
        ini = max(0, ex['inicio'] - 1)
        fim = min(len(seq), ex['fim'])
        for p in range(ini, fim):
            cds += seq[p]; mapa.append(p)
    return cds.upper(), mapa


# ══════════════════════════════════════════════════════════════
# PROCESSAMENTO DOCX
# ══════════════════════════════════════════════════════════════

def limpar_sequencia(doc):
    texto = ''
    for para in doc.paragraphs:
        t = re.sub(r'\d+', '', para.text)
        t = re.sub(r'[\s\r\n\v\f\u000b]+', '', t)
        texto += t
    return re.sub(r'[^ATCGatcg]', '', texto).upper()

def parsear_ncbi_texto(texto):
    exons = []
    p_pos = re.compile(r'exon\s+(\d+)\.\.(\d+)', re.IGNORECASE)
    p_num = re.compile(r'/number=(\d+)')
    mp = list(p_pos.finditer(texto))
    mn = list(p_num.finditer(texto))
    for i, m in enumerate(mp):
        num = int(mn[i].group(1)) if i < len(mn) else i + 1
        exons.append({'numero': num, 'inicio': int(m.group(1)), 'fim': int(m.group(2))})
    return sorted(exons, key=lambda x: x['numero'])

def cor_hex_rgb(h):
    h = h.lstrip('#')
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

def aplicar_run(run, negrito, fonte, tam, cor):
    r, g, b = cor_hex_rgb(cor)
    run.bold = negrito
    run.font.name = fonte
    run.font.size = Pt(tam)
    run.font.color.rgb = RGBColor(r, g, b)

def doc_base():
    doc = Document()
    s = doc.sections[0]
    s.top_margin = s.bottom_margin = Pt(36)
    s.left_margin = s.right_margin = Pt(54)
    return doc

def gerar_completo(seq, cfg, exons, path, chars=60, caixa='maiuscula',
                   codons_grifar=None, utr_regiao=None):
    """
    codons_grifar: lista de dicts [{pos_ini(0-based), pos_fim, cor, negrito}]
    utr_regiao:    dict {pos_ini(0-based), pos_fim, cor} para 5' UTR do éxon 1
    """
    seq = seq.lower() if caixa == 'minuscula' else seq.upper()
    doc = doc_base()
    n = len(seq)

    # Mapa base: None = íntron, exon_obj = éxon
    anot = [None] * n
    for ex in exons:
        for i in range(max(0, ex['inicio']-1), min(n, ex['fim'])):
            anot[i] = ex

    # Mapa UTR: marca região 5' UTR antes do ATG (sobrescreve éxon, fica abaixo do grifo)
    utr_map = [False] * n
    if utr_regiao:
        for i in range(utr_regiao['pos_ini'], min(n, utr_regiao['pos_fim'])):
            utr_map[i] = True

    # Mapa de grifo: sobrescreve tudo
    grifo = [None] * n
    if codons_grifar:
        for g in codons_grifar:
            for i in range(g['pos_ini'], min(n, g['pos_fim'])):
                grifo[i] = g

    for li in range(0, n, chars):
        tr  = seq[li:li+chars]
        al  = anot[li:li+chars]
        gr  = grifo[li:li+chars]
        utr = utr_map[li:li+chars]
        p = doc.add_paragraph()
        p.paragraph_format.space_after = p.paragraph_format.space_before = Pt(0)
        p.style.font.name = cfg['fonte']
        i = 0
        while i < len(tr):
            cur_anot  = al[i]
            cur_grifo = gr[i]
            cur_utr   = utr[i]
            j = i + 1
            while (j < len(tr) and al[j] is cur_anot
                   and gr[j] is cur_grifo and utr[j] == cur_utr):
                j += 1
            run = p.add_run(tr[i:j])
            if cur_grifo is not None:
                # Grifo do códon — prioridade máxima
                aplicar_run(run, cur_grifo.get('negrito', True),
                            cfg['fonte'], cfg.get('tamanho', 11), cur_grifo['cor'])
            elif cur_utr:
                # 5' UTR — cinza escuro, negrito
                cor_utr = utr_regiao.get('cor', '#555555') if utr_regiao else '#555555'
                aplicar_run(run, True, cfg['fonte'], cfg.get('tamanho', 11), cor_utr)
            elif cur_anot is None:
                # Íntron
                aplicar_run(run, False, cfg['fonte'], cfg['tamanho'], cfg['cor'])
            else:
                # Éxon codificante
                aplicar_run(run, True, cur_anot['fonte'], cur_anot['tamanho'], cur_anot['cor'])
            i = j
    doc.save(path)
    return n

def gerar_so_exons(seq, exons, path, chars=60, caixa='maiuscula',
                   codons_grifar=None, utr_regiao=None):
    """
    Gera documento com apenas os éxons concatenados.
    Suporta grifo de códon e marcação de 5' UTR.
    """
    seq_orig = seq.lower() if caixa == 'minuscula' else seq.upper()
    exons_ord = sorted(exons, key=lambda x: x['inicio'])

    # Constrói seq_ex (só éxons) e mapa de posições originais
    seq_ex = ''; mapa_pos = []; mapa_fmt = []
    for idx, ex in enumerate(exons_ord):
        cor = '#000000' if idx % 2 == 0 else '#e67e00'
        obj = {'cor': cor, 'fonte': ex.get('fonte','Courier New'),
               'tamanho': ex.get('tamanho', 11)}
        ini = max(0, ex['inicio']-1); fim = min(len(seq_orig), ex['fim'])
        for p in range(ini, fim):
            seq_ex   += seq_orig[p]
            mapa_pos.append(p)      # posição na seq original (0-based)
            mapa_fmt.append(obj)

    # Mapa de grifo por posição original
    grifo_map = {}
    if codons_grifar:
        for g in codons_grifar:
            for p in range(g['pos_ini'], g['pos_fim']):
                grifo_map[p] = g

    # Mapa UTR por posição original
    utr_set = set()
    if utr_regiao:
        for p in range(utr_regiao['pos_ini'], utr_regiao['pos_fim']):
            utr_set.add(p)

    doc = doc_base()
    total = len(seq_ex)
    for li in range(0, total, chars):
        fatia = seq_ex[li:li+chars]
        pos_l = mapa_pos[li:li+chars]
        fmt_l = mapa_fmt[li:li+chars]
        p = doc.add_paragraph()
        p.paragraph_format.space_after = p.paragraph_format.space_before = Pt(0)
        i = 0
        while i < len(fatia):
            orig_p    = pos_l[i]
            cur_fmt   = fmt_l[i]
            cur_grifo = grifo_map.get(orig_p)
            cur_utr   = orig_p in utr_set
            j = i + 1
            while j < len(fatia):
                g2 = grifo_map.get(pos_l[j])
                if fmt_l[j] is cur_fmt and g2 is cur_grifo and (pos_l[j] in utr_set) == cur_utr:
                    j += 1
                else:
                    break
            run = doc.paragraphs[-1].add_run(fatia[i:j])
            if cur_grifo is not None:
                aplicar_run(run, True, cur_fmt['fonte'], cur_fmt['tamanho'], cur_grifo['cor'])
            elif cur_utr and utr_regiao:
                aplicar_run(run, True, cur_fmt['fonte'], cur_fmt['tamanho'],
                            utr_regiao.get('cor','#555555'))
            else:
                aplicar_run(run, True, cur_fmt['fonte'], cur_fmt['tamanho'], cur_fmt['cor'])
            i = j
    doc.save(path)
    return total

def localizar_stop_codon(seq, exons, proteina_nascente='', coords_cds=None, offset=0):
    """
    Localiza o códon de parada logo após o último aminoácido da proteína nascente.
    Retorna posição no gene, códon (TAA/TAG/TGA) e aminoácidos finais para confirmação.
    """
    STOP_CODONS = {'TAA', 'TAG', 'TGA'}

    # Constrói mapa de posições da CDS
    if coords_cds and offset >= 0:
        mapa_cds = []
        for ini_bloco, fim_bloco in coords_cds:
            ini_adj = ini_bloco - offset
            fim_adj = fim_bloco - offset
            for p in range(max(1, ini_adj), min(len(seq)+1, fim_adj+1)):
                mapa_cds.append(p)
    else:
        cds_raw, mapa_cds = construir_cds(seq, exons)
        mapa_cds = [p+1 for p in mapa_cds]  # converter para 1-based

    if not mapa_cds:
        return {'erro': 'Não foi possível construir mapa da CDS.'}

    n_aa = len(proteina_nascente) if proteina_nascente else len(mapa_cds) // 3

    # Posição do stop codon = logo após o último códon da proteína
    idx_stop_nt = n_aa * 3   # índice no mapa_cds (0-based)

    if idx_stop_nt >= len(mapa_cds):
        return {'erro': f'CDS mapeada ({len(mapa_cds)} nt) menor que esperado ({idx_stop_nt+3} nt).'}

    # Posições genômicas do stop codon
    pos_stop = [mapa_cds[idx_stop_nt + k] for k in range(3)
                if idx_stop_nt + k < len(mapa_cds)]

    if len(pos_stop) < 3:
        return {'erro': 'Stop codon parcialmente fora da sequência mapeada.'}

    # Extrair o códon da sequência
    codon_stop = ''.join(seq[p-1].upper() for p in pos_stop)

    if codon_stop not in STOP_CODONS:
        # Pode ser que a CDS não inclua o stop — tenta as 3 posições seguintes na seq
        p_nt = pos_stop[0] - 1  # 0-based na seq
        codon_stop = seq[p_nt:p_nt+3].upper() if p_nt+3 <= len(seq) else '???'
        if codon_stop not in STOP_CODONS:
            return {
                'erro': None, 'encontrado': False,
                'codon': codon_stop,
                'pos_inicio': pos_stop[0], 'pos_fim': pos_stop[-1],
                'aviso': f'Códon encontrado ({codon_stop}) não é stop codon padrão.'
            }

    # Últimos 3 aminoácidos para confirmação
    ultimos_aas = []
    if proteina_nascente and len(proteina_nascente) >= 3:
        for c in proteina_nascente[-3:]:
            aa = AA_1_TO_3.get(c.upper(), c.upper())
            ultimos_aas.append(aa)

    return {
        'erro': None, 'encontrado': True,
        'codon': codon_stop,
        'pos_inicio': pos_stop[0],
        'pos_fim':    pos_stop[2],
        'pos_gene':   pos_stop,
        'ultimos_aas': ultimos_aas,
        'n_aa_total':  n_aa,
    }


# Mapa 1-letra → 3-letras para conversão da proteína
AA_1_TO_3 = {
    'A':'Ala','R':'Arg','N':'Asn','D':'Asp','C':'Cys','Q':'Gln','E':'Glu',
    'G':'Gly','H':'His','I':'Ile','L':'Leu','K':'Lys','M':'Met','F':'Phe',
    'P':'Pro','S':'Ser','T':'Thr','W':'Trp','Y':'Tyr','V':'Val','*':'Stop',
}

def analisar_codons(seq, exons, aas_entrada, proteina_nascente='', coords_cds=None, offset=0):
    """
    Analisa os primeiros códons.
    Se proteina_nascente disponível (importada do GenBank via /translation):
      → compara aa a aa diretamente — ignora problemas de UTR e offset
    Caso contrário:
      → tenta construir CDS a partir dos éxons (menos confiável)
    """
    # Normaliza aminoácidos de entrada
    aas = []
    for a in aas_entrada:
        norm = normalizar_aa(a)
        if not norm: return {'erro': f"Aminoácido não reconhecido: '{a}'"}
        aas.append(norm)
    if len(aas) < 3: return {'erro': 'Forneça no mínimo 3 aminoácidos.'}
    n_aa = len(aas)

    # ── Modo 1: usar proteína nascente diretamente (mais confiável) ──
    if proteina_nascente and len(proteina_nascente) >= n_aa:
        # Converte proteína de 1-letra para 3-letras
        prot_3 = [AA_1_TO_3.get(c.upper(), '???') for c in proteina_nascente[:n_aa]]
        matches = [aas[i] == prot_3[i] for i in range(n_aa)]
        todos_ok = all(matches)

        # Calcular posições no gene usando coords_cds ajustadas pelo offset
        pos_gene = []
        if coords_cds and offset >= 0:
            # Reconstrói mapa de posições da CDS no gene
            mapa_cds = []
            for ini_bloco, fim_bloco in coords_cds:
                ini_adj = ini_bloco - offset
                fim_adj = fim_bloco - offset
                for p in range(max(1, ini_adj), min(len(seq)+1, fim_adj+1)):
                    mapa_cds.append(p)
            for i in range(n_aa):
                idx = i * 3
                if idx < len(mapa_cds):
                    pos_gene.append(mapa_cds[idx])
                else:
                    pos_gene.append(0)
        else:
            pos_gene = [0] * n_aa

        # Códons reais: extrair da sequência usando as posições calculadas
        codons_reais = []
        for i in range(n_aa):
            p = pos_gene[i] - 1 if pos_gene[i] > 0 else 0
            if p+3 <= len(seq):
                codons_reais.append(seq[p:p+3].upper())
            else:
                codons_reais.append('???')

        if todos_ok:
            return {'erro':None,'encontrado':True,'aas_fornecidos':aas,
                    'combinacao':codons_reais,'codons_na_cds':codons_reais,
                    'posicoes_gene':pos_gene,
                    'pos_gene_inicio':pos_gene[0] if pos_gene else 0,
                    'pos_gene_fim':pos_gene[-1]+2 if pos_gene else 0,
                    'matches':matches,'fonte':'translation'}
        return {'erro':None,'encontrado':False,'aas_fornecidos':aas,
                'aas_na_cds':prot_3,'codons_na_cds':codons_reais,
                'posicoes_gene':pos_gene,'matches':matches,'fonte':'translation'}

    # ── Modo 2: construir CDS a partir dos éxons (fallback) ──
    cds, mapa = construir_cds(seq, exons)
    if len(cds) < n_aa*3:
        return {'erro': 'CDS muito curta. Se importou do NCBI, verifique se os éxons foram carregados.'}
    codons_cds = [cds[i*3:(i+1)*3].upper() for i in range(n_aa)]
    aas_cds    = [CODIGO_GENETICO.get(c,'???') for c in codons_cds]
    pos_gene   = [mapa[i*3] for i in range(n_aa)]
    matches    = [aas[i] == aas_cds[i] for i in range(n_aa)]
    todos_ok   = all(matches)

    if todos_ok:
        return {'erro':None,'encontrado':True,'aas_fornecidos':aas,
                'combinacao':codons_cds,'codons_na_cds':codons_cds,
                'posicoes_gene':pos_gene,'pos_gene_inicio':mapa[0],
                'pos_gene_fim':mapa[n_aa*3-1],'matches':matches,'fonte':'exons'}
    return {'erro':None,'encontrado':False,'aas_fornecidos':aas,
            'aas_na_cds':aas_cds,'codons_na_cds':codons_cds,
            'posicoes_gene':pos_gene,'matches':matches,'fonte':'exons'}

def localizar_codon(seq, exons, numero_aa):
    if numero_aa < 1: return {'erro': 'Número deve ser ≥ 1.'}
    cds, mapa = construir_cds(seq, exons)
    i0 = (numero_aa-1)*3; i1 = numero_aa*3
    if i1 > len(cds):
        return {'erro': f"Aminoácido {numero_aa} além da CDS ({len(cds)//3} aa)."}
    codon = cds[i0:i1].upper()
    return {'erro':None,'numero_aa':numero_aa,'codon':codon,
            'aminoacido':CODIGO_GENETICO.get(codon,'???'),
            'pos_gene':[mapa[i0+k] for k in range(3)],
            'pos_inicio':mapa[i0]+1,'pos_fim':mapa[i1-1]+1}


# ══════════════════════════════════════════════════════════════
# NCBI API
# ══════════════════════════════════════════════════════════════

ENTREZ = 'https://eutils.ncbi.nlm.nih.gov/entrez/eutils'

def _get(url, timeout=25):
    for t in range(3):
        try:
            with urllib.request.urlopen(url, timeout=timeout) as r:
                return r.read().decode('utf-8', errors='replace')
        except Exception as e:
            if t == 2: raise RuntimeError(f"Erro NCBI: {e}")
            time.sleep(1.5)

def ncbi_buscar(termo, email, retmax=8):
    p = urllib.parse.urlencode({'db':'nuccore','term':termo,'retmax':retmax,
                                'retmode':'json','email':email,'tool':'ExonEditor'})
    dados = json.loads(_get(f"{ENTREZ}/esearch.fcgi?{p}"))
    ids = dados.get('esearchresult',{}).get('idlist',[])
    if not ids: return []
    p2 = urllib.parse.urlencode({'db':'nuccore','id':','.join(ids),
                                 'retmode':'json','email':email,'tool':'ExonEditor'})
    summ = json.loads(_get(f"{ENTREZ}/esummary.fcgi?{p2}"))
    result = summ.get('result',{})
    uids = result.get('uids', ids)
    regs = []
    for uid in uids:
        info = result.get(uid,{})
        regs.append({'id':uid,'acc':info.get('accessionversion',uid),
                     'titulo':info.get('title','—'),'len':info.get('slen','?')})
    return regs

def ncbi_fetch_genbank(acc, email, seq_start=None, seq_stop=None):
    params = {'db':'nuccore','id':acc,'rettype':'gb',
              'retmode':'text','email':email,'tool':'ExonEditor'}
    if seq_start and seq_stop:
        params['seq_start'] = seq_start
        params['seq_stop']  = seq_stop
        params['strand']    = 1
    p = urllib.parse.urlencode(params)
    return _get(f"{ENTREZ}/efetch.fcgi?{p}", timeout=35)

def _parsear_genbank_biopython(record):
    """Parser usando Biopython — muito mais confiável."""
    res = {
        'sequencia':       str(record.seq).upper(),
        'exons':           [],
        'proteinas':       [],
        'organismo':       '',
        'definicao':       record.description,
        'accession':       record.id,
        'regiao_sugerida': None,
        'proteina_nascente': '',
        'proteina_madura': [],
    }

    # Organismo
    for f in record.features:
        if f.type == 'source':
            res['organismo'] = f.qualifiers.get('organism',[''])[0]
            break

    # Éxons
    for f in record.features:
        if f.type == 'exon':
            nums = f.qualifiers.get('number', [None])
            num  = int(nums[0]) if nums[0] else len(res['exons'])+1
            res['exons'].append({
                'numero': num,
                'inicio': int(f.location.start) + 1,  # 1-based
                'fim':    int(f.location.end),
            })
    res['exons'].sort(key=lambda x: x['numero'])

    # CDS — proteína nascente
    for f in record.features:
        if f.type == 'CDS':
            trans = f.qualifiers.get('translation',[''])[0]
            prod  = f.qualifiers.get('product',['—'])[0]
            pid   = f.qualifiers.get('protein_id',['—'])[0]
            # Coordenadas dos blocos da CDS
            try:
                coords_cds = [(int(part.start)+1, int(part.end))
                              for part in f.location.parts]
            except Exception:
                coords_cds = []
            res['proteinas'].append({
                'produto':    prod,
                'protein_id': pid,
                'sequencia':  trans,
                'n_aa':       len(trans),
                'isoforma':   '',
                'coords_cds': coords_cds,
            })
            if trans and not res['proteina_nascente']:
                res['proteina_nascente'] = trans

    # mat_peptide — proteína madura
    for f in record.features:
        if f.type == 'mat_peptide':
            prod = f.qualifiers.get('product',['—'])[0]
            try:
                coords = [(int(part.start)+1, int(part.end))
                          for part in f.location.parts]
            except Exception:
                coords = [(int(f.location.start)+1, int(f.location.end))]
            res['proteina_madura'].append({
                'produto': prod,
                'nota':    f.qualifiers.get('note',[''])[0],
                'coords':  coords,
            })

    # Região sugerida
    if res['exons']:
        res['regiao_sugerida'] = {
            'de':  res['exons'][0]['inicio'],
            'ate': res['exons'][-1]['fim'],
        }

    return res


def parsear_genbank(gb):
    """
    Parser robusto do formato GenBank.
    Usa Biopython se disponível, fallback para parser manual.
    """
    # Tenta usar Biopython primeiro (mais robusto)
    try:
        from Bio import SeqIO
        from io import StringIO
        record = SeqIO.read(StringIO(gb), 'genbank')
        return _parsear_genbank_biopython(record)
    except ImportError:
        pass  # Biopython não instalado, usa parser manual
    except Exception as e:
        print(f"[DEBUG] Biopython falhou: {e}, usando parser manual")

    res = {'sequencia':'','exons':[],'proteinas':[],'organismo':'',
           'definicao':'','accession':'','regiao_sugerida':None,
           'proteina_nascente':'','proteina_madura':[]}
    linhas = gb.splitlines()
    for l in linhas:
        if l.startswith('DEFINITION'): res['definicao'] = l[12:].strip()
        elif l.startswith('ACCESSION'):  res['accession'] = l[12:].strip().split()[0]
        elif '  ORGANISM  ' in l:        res['organismo'] = l.strip().replace('ORGANISM','').strip()

    em_ft = False; bloco = ''; tipo = ''
    exons_raw = []; cds_raw = []; gene_raw = []; mrna_raw = []; mat_raw = []
    for l in linhas:
        if l.startswith('FEATURES'):   em_ft = True; continue
        if l.startswith('ORIGIN') or l.startswith('CONTIG'): em_ft = False
        if em_ft:
            if len(l)>5 and l[5]!=' ' and l[:5]=='     ':
                if tipo=='exon':        exons_raw.append(bloco)
                elif tipo=='CDS':       cds_raw.append(bloco)
                elif tipo=='gene':      gene_raw.append(bloco)
                elif tipo=='mRNA':      mrna_raw.append(bloco)
                elif tipo in ('mat_peptide','Mature_peptide'): mat_raw.append(bloco)
                pts = l.strip().split(None,1); tipo = pts[0] if pts else ''
                bloco = l.strip()
            else: bloco += '\n' + l.strip()
    if tipo=='exon':        exons_raw.append(bloco)
    elif tipo=='CDS':       cds_raw.append(bloco)
    elif tipo=='gene':      gene_raw.append(bloco)
    elif tipo=='mRNA':      mrna_raw.append(bloco)
    elif tipo in ('mat_peptide','Mature_peptide'): mat_raw.append(bloco)

    p_coord = re.compile(r'(?:complement\()?(\d+)\.\.(\d+)\)?')
    p_join  = re.compile(r'(\d+)\.\.(\d+)')
    p_num   = re.compile(r'/number=(\d+)')

    for b in exons_raw:
        mc = p_coord.search(b); mn = p_num.search(b)
        if mc:
            num = int(mn.group(1)) if mn else len(res['exons'])+1
            res['exons'].append({'numero':num,'inicio':int(mc.group(1)),'fim':int(mc.group(2))})
    res['exons'].sort(key=lambda x: x['numero'])

    p_prod  = re.compile(r'/product="([^"]+)"')
    p_prot  = re.compile(r'/protein_id="([^"]+)"')
    p_trans = re.compile(r'/translation="([^"]+)"', re.DOTALL)
    p_iso   = re.compile(r'/note="([^"]*isoform[^"]*)"', re.IGNORECASE)
    for b in cds_raw:
        mt = p_trans.search(b)
        seq_p = re.sub(r'\s+','',mt.group(1)) if mt else ''
        mp2 = p_prod.search(b); mi = p_prot.search(b); miso = p_iso.search(b)
        # Extrai coordenadas dos blocos da CDS (para calcular offset do mat_peptide)
        coords_cds = [(int(a),int(b2)) for a,b2 in p_join.findall(b)]
        res['proteinas'].append({'produto':mp2.group(1) if mp2 else '—',
                                  'protein_id':mi.group(1) if mi else '—',
                                  'sequencia':seq_p,'n_aa':len(seq_p),
                                  'isoforma':miso.group(1) if miso else '',
                                  'coords_cds':coords_cds})
        # Primeira CDS com translation = proteína nascente
        if seq_p and not res['proteina_nascente']:
            res['proteina_nascente'] = seq_p

    # ── Proteína madura (mat_peptide) ──────────────────────────────
    p_prod_mat = re.compile(r'/product="([^"]+)"')
    p_note_mat = re.compile(r'/note="([^"]+)"')
    for b in mat_raw:
        mp = p_prod_mat.search(b)
        mn_note = p_note_mat.search(b)
        produto = mp.group(1) if mp else '—'
        nota    = mn_note.group(1) if mn_note else ''
        # mat_peptide não tem /translation — a sequência vem do NCBI separadamente
        # guardamos coordenadas para identificar a proteína madura
        coords_mat = p_join.findall(b)
        res['proteina_madura'].append({
            'produto': produto,
            'nota':    nota,
            'coords':  [(int(a),int(b2)) for a,b2 in coords_mat],
        })

    # ── Região sugerida ──────────────────────────────────────────
    # Equivalente ao "Selected region" do NCBI:
    # do início do 1º éxon ao fim do último éxon (inclui íntrons entre eles)
    # Prioridade: éxons detectados > mRNA span > gene feature
    coords = None
    if res['exons']:
        coords = (res['exons'][0]['inicio'], res['exons'][-1]['fim'])
    elif mrna_raw:
        todos = p_join.findall(mrna_raw[0])
        if todos:
            inis = [int(x[0]) for x in todos]
            fins = [int(x[1]) for x in todos]
            coords = (min(inis), max(fins))
    elif gene_raw:
        mc = p_coord.search(gene_raw[0])
        if mc: coords = (int(mc.group(1)), int(mc.group(2)))

    if coords:
        res['regiao_sugerida'] = {'de': coords[0], 'ate': coords[1]}

    em_orig = False; seq = ''
    for l in linhas:
        if l.startswith('ORIGIN'): em_orig = True; continue
        if em_orig:
            if l.startswith('//'): break
            seq += re.sub(r'[\d\s]','',l)
    res['sequencia'] = seq.upper()
    return res


# ══════════════════════════════════════════════════════════════
# PALETA E CONSTANTES
# ══════════════════════════════════════════════════════════════

C = {
    'fundo':    '#f4f6fb',
    'header':   '#1a3a5c',
    'card':     '#ffffff',
    'borda':    '#d1d9e6',
    'amber':    '#b45309',
    'amber_bg': '#fffbeb',
    'green':    '#15803d',
    'green_bg': '#f0fdf4',
    'azul':     '#2563eb',
    'azul_bg':  '#eff6ff',
    'roxo':     '#7c3aed',
    'roxo_bg':  '#f5f3ff',
    'cinza':    '#9ca3af',
    'texto':    '#1e293b',
    'sub':      '#64748b',
    'laranja':  '#c2410c',
    'laranja_bg':'#fff7ed',
}
FONTES = ['Courier New','Consolas','Lucida Console','Monaco','Arial','Times New Roman']


# ══════════════════════════════════════════════════════════════
# MIXIN — EDITOR COMPARTILHADO (éxons, barra, códons, documentos)
# ══════════════════════════════════════════════════════════════

class EditorMixin:
    """
    Mixin com todos os métodos compartilhados entre os dois modos.
    Requer que a subclasse defina self.exons, self.sequencia_verificada,
    self.codon_marcado e os widgets de formatação.
    """

    # ── helpers UI ───────────────────────────────

    def _tooltip(self, widget, texto):
        """Exibe tooltip ao passar o mouse sobre widget."""
        def _show(e):
            tip = tk.Toplevel(widget)
            tip.wm_overrideredirect(True)
            tip.attributes('-topmost', True)
            x = widget.winfo_rootx() + 20
            y = widget.winfo_rooty() + 24
            tip.geometry(f'+{x}+{y}')
            fr = tk.Frame(tip, bg='#1e293b',
                          highlightbackground='#475569', highlightthickness=1)
            fr.pack()
            tk.Label(fr, text=texto, font=('Arial',9), bg='#1e293b', fg='white',
                     padx=10, pady=8, justify=tk.LEFT, wraplength=320).pack()
            widget._tip = tip
            def _hide(e2=None):
                try: tip.destroy()
                except Exception: pass
            tip.bind('<Leave>', _hide)
            widget.bind('<Leave>', _hide)
        widget.bind('<Enter>', _show)
        return widget

    def _btn_help(self, parent, texto, bg=C['card']):
        """Cria label '?' que exibe tooltip ao passar o mouse."""
        btn = tk.Label(parent, text=' ？ ', font=('Arial',8,'bold'),
                       bg=bg, fg='#94a3b8',
                       relief=tk.FLAT, cursor='question_arrow')
        self._tooltip(btn, texto)
        return btn

    def _card(self, parent, titulo, cor_t=None, bg=None, pady_top=8):
        bg = bg or C['card']
        fr = tk.Frame(parent, bg=bg, highlightbackground=C['borda'], highlightthickness=1)
        fr.pack(fill=tk.X, pady=(pady_top,0))
        if titulo:
            tk.Label(fr, text=titulo, font=('Arial',9,'bold'), bg=bg,
                     fg=cor_t or C['header'], padx=12, pady=5).pack(anchor=tk.W)
            tk.Frame(fr, bg=C['borda'], height=1).pack(fill=tk.X)
        inner = tk.Frame(fr, bg=bg, padx=12, pady=8)
        inner.pack(fill=tk.X)
        return inner

    def _btn(self, parent, text, cmd, bg, fg='white', size=9,
             state=tk.NORMAL, padx=12, pady=5, side=tk.LEFT, anchor=None):
        b = tk.Button(parent, text=text, command=cmd, bg=bg, fg=fg,
                      font=('Arial',size,'bold'), relief=tk.FLAT,
                      padx=padx, pady=pady, state=state, cursor='hand2')
        if anchor: b.pack(anchor=anchor, pady=2)
        else:      b.pack(side=side, padx=(0,6))
        return b

    # ── Formatação base ──────────────────────────

    def _build_fmt_base(self, parent):
        c = self._card(parent, '② Formatação Base  (Íntrons)')
        fr = tk.Frame(c, bg=C['card']); fr.pack(fill=tk.X)
        btn_h = self._btn_help(fr, 'Define a fonte e o tamanho dos íntrons no documento Word.\n'
                               'Cor cinza é o padrão para íntrons.\n'
                               'Caixa: escolha entre MAIÚSCULA (ATCG) ou minúscula (atcg).')
        btn_h.pack(side=tk.RIGHT, padx=(4,0))
        tk.Label(fr, text='Fonte:', bg=C['card'], font=('Arial',9), fg=C['sub']).pack(side=tk.LEFT)
        ttk.Combobox(fr, textvariable=self.fonte_base, values=FONTES,
                     width=15, state='readonly').pack(side=tk.LEFT, padx=(4,12))
        tk.Label(fr, text='Tamanho:', bg=C['card'], font=('Arial',9), fg=C['sub']).pack(side=tk.LEFT)
        tk.Spinbox(fr, from_=6, to=24, textvariable=self.tamanho_base,
                   width=4, font=('Arial',9)).pack(side=tk.LEFT, padx=(4,12))
        tk.Label(fr, text='Cor:', bg=C['card'], font=('Arial',9), fg=C['sub']).pack(side=tk.LEFT)
        self.btn_cor_base = tk.Button(fr, text='  ██  ', bg=self.cor_base, fg=self.cor_base,
                                      command=self._escolher_cor_base, relief=tk.RAISED, padx=4)
        self.btn_cor_base.pack(side=tk.LEFT, padx=(4,6))
        tk.Label(fr, text='(cinza — padrão íntrons)', font=('Arial',8,'italic'),
                 bg=C['card'], fg=C['cinza']).pack(side=tk.LEFT)
        fr2 = tk.Frame(c, bg=C['card']); fr2.pack(fill=tk.X, pady=(6,0))
        tk.Label(fr2, text='Caixa:', bg=C['card'], font=('Arial',9), fg=C['sub']).pack(side=tk.LEFT)
        for val, txt in [('maiuscula','MAIÚSCULA  (ATCG)'),('minuscula','minúscula  (atcg)')]:
            tk.Radiobutton(fr2, text=txt, variable=self.caixa, value=val,
                           bg=C['card'], font=('Arial',9),
                           activebackground=C['card']).pack(side=tk.LEFT, padx=(8,0))

    def _escolher_cor_base(self):
        cor = colorchooser.askcolor(color=self.cor_base, title='Cor Base')[1]
        if cor: self.cor_base = cor; self.btn_cor_base.configure(bg=cor, fg=cor)

    def _escolher_cor_exon(self):
        cor = colorchooser.askcolor(color=self.cor_exon_atual, title='Cor do Éxon')[1]
        if cor: self.cor_exon_atual = cor; self.btn_cor_exon.configure(bg=cor, fg=cor)

    # ── Tabela de éxons ──────────────────────────

    def _build_tabela_exons(self, parent):
        # Controles de adição manual
        c = self._card(parent, '③ Intervalos de Éxons')
        self._btn_help(c, 'Defina os intervalos de cada éxon (posição De e Até na sequência).\n'
                   'Você pode importar automaticamente colando o texto do GenBank,\n'
                   'ou adicionar manualmente. Os éxons são numerados na ordem de adição.').pack(anchor=tk.E)
        self.lbl_bloqueio = tk.Label(c,
            text='⚠  Confirme a sequência para habilitar.',
            font=('Arial',8,'italic'), bg=C['card'], fg=C['amber'])
        self.lbl_bloqueio.pack(anchor=tk.W, pady=(0,6))

        # NCBI texto (importador manual de texto colado)
        fr_ncbi_txt = tk.Frame(c, bg=C['card']); fr_ncbi_txt.pack(fill=tk.X, pady=(0,8))
        self.btn_ncbi_txt = self._btn(fr_ncbi_txt, '📋  Importar texto do NCBI',
                                      self._abrir_importador_ncbi_txt, C['roxo'],
                                      size=9, padx=10, pady=4, state=tk.DISABLED)
        tk.Label(fr_ncbi_txt, text='Cole o texto da seção Features do GenBank',
                 font=('Arial',8,'italic'), bg=C['card'], fg=C['cinza']).pack(side=tk.LEFT)

        # Manual
        fr1 = tk.Frame(c, bg=C['card']); fr1.pack(fill=tk.X)
        tk.Label(fr1, text='Início:', bg=C['card'], font=('Arial',9), fg=C['sub']).pack(side=tk.LEFT)
        self.entry_inicio = tk.Entry(fr1, width=8, font=('Courier New',10),
                                     state=tk.DISABLED, relief=tk.SOLID, bd=1, bg='#f8fafc')
        self.entry_inicio.pack(side=tk.LEFT, padx=(3,10), ipady=3)
        tk.Label(fr1, text='Fim:', bg=C['card'], font=('Arial',9), fg=C['sub']).pack(side=tk.LEFT)
        self.entry_fim = tk.Entry(fr1, width=8, font=('Courier New',10),
                                  state=tk.DISABLED, relief=tk.SOLID, bd=1, bg='#f8fafc')
        self.entry_fim.pack(side=tk.LEFT, padx=(3,12), ipady=3)

        fr2 = tk.Frame(c, bg=C['card']); fr2.pack(fill=tk.X, pady=(6,0))
        tk.Label(fr2, text='Fonte:', bg=C['card'], font=('Arial',9), fg=C['sub']).pack(side=tk.LEFT)
        self.combo_fonte_exon = ttk.Combobox(fr2, textvariable=self.fonte_exon,
                                              values=FONTES[:4], width=13, state=tk.DISABLED)
        self.combo_fonte_exon.pack(side=tk.LEFT, padx=(3,10))
        tk.Label(fr2, text='Tam:', bg=C['card'], font=('Arial',9), fg=C['sub']).pack(side=tk.LEFT)
        self.spin_tam = tk.Spinbox(fr2, from_=6, to=24, textvariable=self.tamanho_exon,
                                   width=4, font=('Arial',9), state=tk.DISABLED)
        self.spin_tam.pack(side=tk.LEFT, padx=(3,10))
        tk.Label(fr2, text='Cor:', bg=C['card'], font=('Arial',9), fg=C['sub']).pack(side=tk.LEFT)
        self.btn_cor_exon = tk.Button(fr2, text='  ██  ', bg=self.cor_exon_atual,
                                      fg=self.cor_exon_atual, command=self._escolher_cor_exon,
                                      relief=tk.RAISED, padx=3, state=tk.DISABLED)
        self.btn_cor_exon.pack(side=tk.LEFT, padx=(3,4))
        tk.Label(fr2, text='(preto)', font=('Arial',7,'italic'),
                 bg=C['card'], fg=C['cinza']).pack(side=tk.LEFT, padx=(0,8))
        self.btn_add = self._btn(fr2, '＋ Adicionar',
                                 self._adicionar_exon, C['cinza'],
                                 size=8, padx=10, pady=4, state=tk.DISABLED)

        # Tabela
        fr_tab = tk.Frame(parent, bg=C['fundo']); fr_tab.pack(fill=tk.X, pady=(6,0))
        cols = ('Éxon','Início','Fim','Tamanho','Fonte','Cor')
        self.tabela = ttk.Treeview(fr_tab, columns=cols, show='headings', height=5)
        larg = {'Éxon':65,'Início':75,'Fim':75,'Tamanho':85,'Fonte':140,'Cor':75}
        for col in cols:
            self.tabela.heading(col, text=col)
            self.tabela.column(col, width=larg.get(col,80), anchor=tk.CENTER)
        sc = ttk.Scrollbar(fr_tab, orient=tk.VERTICAL, command=self.tabela.yview)
        self.tabela.configure(yscrollcommand=sc.set)
        self.tabela.pack(side=tk.LEFT, fill=tk.X, expand=True)
        sc.pack(side=tk.RIGHT, fill=tk.Y)

        fr_ctrl = tk.Frame(parent, bg=C['fundo']); fr_ctrl.pack(fill=tk.X, pady=(4,0))
        self.btn_rem   = self._btn(fr_ctrl,'🗑  Remover',self._remover_exon,'#dc2626',
                                   size=8,padx=8,pady=3,state=tk.DISABLED)
        self.btn_clear = self._btn(fr_ctrl,'🗑  Limpar Todos',self._limpar_exons,C['cinza'],
                                   size=8,padx=8,pady=3,state=tk.DISABLED)
        self.lbl_cont  = tk.Label(fr_ctrl, text='', font=('Arial',8,'italic'),
                                  bg=C['fundo'], fg=C['green'])
        self.lbl_cont.pack(side=tk.LEFT, padx=8)

    # ── Barra visual ─────────────────────────────

    def _build_barra(self, parent):
        c = self._card(parent, '🗺  Mapa do Gene  —  Éxons (■) e Íntrons (□)', cor_t=C['header'])
        self._btn_help(c, 'Representação proporcional do gene.\n'
                   'Éxons (■) em preto, íntrons (□) em cinza.\n'
                   'As setas coloridas marcam códons localizados:\n'
                   '  🔴 Vermelho = Stop codon\n'
                   '  🟣 Roxo = Códon localizado por nº de aminoácido').pack(anchor=tk.E)
        self.canvas = tk.Canvas(c, height=44, bg='white',
                                highlightthickness=1, highlightbackground=C['borda'])
        self.canvas.pack(fill=tk.X)
        self.canvas.bind('<Configure>', lambda e: self._draw())
        fr_barra_bot = tk.Frame(c, bg=C['card']); fr_barra_bot.pack(fill=tk.X, pady=(4,0))
        self.lbl_barra = tk.Label(fr_barra_bot, text='Aguardando éxons…',
                                   font=('Arial',8,'italic'), bg=C['card'], fg=C['sub'])
        self.lbl_barra.pack(side=tk.LEFT)
        tk.Button(fr_barra_bot, text='💾  Salvar imagem',
                  command=self._salvar_barra,
                  bg=C['header'], fg='white', font=('Arial',8,'bold'),
                  relief=tk.FLAT, padx=8, pady=2,
                  cursor='hand2').pack(side=tk.RIGHT)

    def _voltar(self):
        try: self.inner.unbind_all('<MouseWheel>')
        except Exception: pass
        self.destroy()
        self.master.deiconify()

    def _salvar_barra(self):
        try:
            from PIL import ImageGrab
        except ImportError:
            messagebox.showerror('Erro',
                'Instale a biblioteca Pillow para salvar imagens:\n\npip install Pillow')
            return
        path = filedialog.asksaveasfilename(
            defaultextension='.png',
            filetypes=[('PNG','*.png'),('JPEG','*.jpg')],
            title='Salvar mapa do gene como…')
        if not path: return
        try:
            cv = self.canvas
            # Coordenadas absolutas do canvas na tela
            x = cv.winfo_rootx(); y = cv.winfo_rooty()
            w = cv.winfo_width(); h = cv.winfo_height()
            img = ImageGrab.grab(bbox=(x, y, x+w, y+h))
            img.save(path)
            messagebox.showinfo('Salvo', f'Imagem salva em:\n{path}')
        except Exception as e:
            messagebox.showerror('Erro ao salvar', str(e))

    def _draw(self):
        cv = self.canvas; cv.delete('all')
        total = len(self.sequencia_verificada) if self.sequencia_verificada else 0
        if total == 0: return
        W = cv.winfo_width(); H = cv.winfo_height()
        if W <= 1: return
        # Fundo cinza uniforme = íntrons (mesma altura que os éxons)
        cv.create_rectangle(0, 4, W, H-4, fill='#aaaaaa', outline='')
        pb = 0
        for ex in self.exons:
            ini = ex['inicio']-1; fim = ex['fim']
            x1 = int((ini/total)*W); x2 = max(x1+2, int((fim/total)*W))
            # Éxons mesma altura que íntrons (0..H), cor definida (preto padrão)
            cv.create_rectangle(x1, 4, x2, H-4, fill=ex.get('cor','#000000'), outline='')
            pb += (fim-ini)

        # Desenha marcadores (painel B=roxo, painel C=vermelho)
        marcadores = [m for m in [
            getattr(self, 'marc_stop',  None),
            getattr(self, 'marc_loc',   None),
        ] if m is not None]
        for cm in marcadores:
            xc = int(((cm['pi'] + cm['pf']) / 2 / total) * W)
            xc = max(8, min(W-8, xc))
            cs = cm['cor']
            cv.create_line(xc, 0, xc, H, fill=cs, width=2, dash=(4,2))
            cv.create_polygon(xc-7, 2, xc+7, 2, xc, 14, fill=cs, outline='')
            cv.create_text(xc, H-2, text=cm['label'],
                           font=('Arial',7,'bold'), fill=cs, anchor='s')

        pct = pb/total*100 if total > 0 else 0
        n   = len(self.exons)
        ex_txts = [f'🎯 {m["label"]}: {m["pi"]+1}–{m["pf"]+1} pb' for m in marcadores]
        ex_txt  = '  |  ' + '   '.join(ex_txts) if ex_txts else ''
        self.lbl_barra.config(
            text=(f'{n} éxon(s) — {pb:,} pb codificantes ({pct:.1f}% do gene){ex_txt}'
                  if n > 0 else 'Nenhum éxon adicionado ainda.'))

    # ── Análise de códons ────────────────────────

    def _build_codons(self, parent):
        c = self._card(parent, '④ Análise de Códons e Proteínas', cor_t=C['header'])
        self.lbl_bl_cod = tk.Label(c, text='⚠  Confirme a sequência e adicione Éxons.',
                                    font=('Arial',8,'italic'), bg=C['card'], fg=C['amber'])
        self.lbl_bl_cod.pack(anchor=tk.W, pady=(0,6))

        # ── Sub-painel: Sequências de Proteínas (importadas do GenBank) ──
        fr_prot = tk.Frame(c, bg='#faf5ff',
                           highlightbackground='#d8b4fe', highlightthickness=1,
                           padx=10, pady=8)
        fr_prot.pack(fill=tk.X, pady=(0,10))
        fr_prot_tit = tk.Frame(fr_prot, bg='#faf5ff'); fr_prot_tit.pack(fill=tk.X)
        tk.Label(fr_prot_tit, text='🧪  Sequências Proteicas  (importadas automaticamente do GenBank)',
                 font=('Arial',9,'bold'), bg='#faf5ff', fg='#7c3aed').pack(side=tk.LEFT)
        self._btn_help(fr_prot_tit,
            'Proteína Nascente (preproprotein):\n'
            '  Sequência completa traduzida da CDS, incluindo\n'
            '  peptídeo sinal e propeptídeo.\n\n'
            'Proteína Madura (Mature Peptide):\n'
            '  Forma ativa após clivagem proteolítica.\n'
            '  Calculada a partir das coordenadas do mat_peptide\n'
            '  no GenBank.',
            bg='#faf5ff').pack(side=tk.RIGHT)

        # Proteína Nascente
        tk.Label(fr_prot, text='Proteína Nascente (preproprotein):',
                 font=('Arial',8,'bold'), bg='#faf5ff', fg=C['texto']).pack(anchor=tk.W, pady=(6,2))
        fr_nasc = tk.Frame(fr_prot, bg='#faf5ff'); fr_nasc.pack(fill=tk.X)
        self.txt_prot_nasc = tk.Text(fr_nasc, height=3, font=('Courier New',8),
                                      bg='#f3e8ff', relief=tk.FLAT, bd=1,
                                      state=tk.DISABLED, wrap=tk.WORD)
        sc_nasc = ttk.Scrollbar(fr_nasc, command=self.txt_prot_nasc.yview)
        self.txt_prot_nasc.configure(yscrollcommand=sc_nasc.set)
        self.txt_prot_nasc.pack(side=tk.LEFT, fill=tk.X, expand=True)
        sc_nasc.pack(side=tk.RIGHT, fill=tk.Y)

        # Proteína Madura
        tk.Label(fr_prot, text='Proteína Madura (Mature Peptide):',
                 font=('Arial',8,'bold'), bg='#faf5ff', fg=C['texto']).pack(anchor=tk.W, pady=(8,2))
        fr_mad = tk.Frame(fr_prot, bg='#faf5ff'); fr_mad.pack(fill=tk.X)
        self.txt_prot_mad = tk.Text(fr_mad, height=3, font=('Courier New',8),
                                     bg='#ede9fe', relief=tk.FLAT, bd=1,
                                     state=tk.DISABLED, wrap=tk.WORD)
        sc_mad = ttk.Scrollbar(fr_mad, command=self.txt_prot_mad.yview)
        self.txt_prot_mad.configure(yscrollcommand=sc_mad.set)
        self.txt_prot_mad.pack(side=tk.LEFT, fill=tk.X, expand=True)
        sc_mad.pack(side=tk.RIGHT, fill=tk.Y)

        self.lbl_prot_info = tk.Label(fr_prot, text='Aguardando importação do GenBank…',
                                       font=('Arial',7,'italic'), bg='#faf5ff', fg=C['cinza'])
        self.lbl_prot_info.pack(anchor=tk.W, pady=(4,0))

        fr_cod = tk.Frame(c, bg=C['card']); fr_cod.pack(fill=tk.X)
        fr_cod.columnconfigure(0, weight=1); fr_cod.columnconfigure(1, weight=1)
        fr_cod.columnconfigure(2, weight=1)

        # Painel A
        fa = tk.Frame(fr_cod, bg=C['green_bg'],
                      highlightbackground='#86efac', highlightthickness=1,
                      padx=10, pady=8)
        fa.grid(row=0, column=0, sticky='nsew', padx=(0,6))
        fr_a_tit = tk.Frame(fa, bg=C['green_bg']); fr_a_tit.pack(fill=tk.X)
        tk.Label(fr_a_tit, text='🔬  Primeiros códons',
                 font=('Arial',9,'bold'), bg=C['green_bg'], fg=C['green']).pack(side=tk.LEFT)
        self._btn_help(fr_a_tit,
            'Verifica se os primeiros aminoácidos fornecidos\n'
            'correspondem ao início da proteína nascente (CDS).\n\n'
            '⚡ Auto (Preprotein): usa automaticamente os primeiros\n'
            '6 aminoácidos da proteína nascente importada do GenBank.\n\n'
            '☐ Isolar 5\' UTR: destaca em azul os nucleotídeos do\n'
            'Éxon 1 antes do ATG no documento Word.',
            bg=C['green_bg']).pack(side=tk.RIGHT)
        tk.Label(fa, text='≥3 aminoácidos (1 ou 3 letras, vírgula/espaço)\nEx: Met,Lys,Ser  ou  M K S',
                 font=('Arial',8), bg=C['green_bg'], fg=C['sub'],
                 justify=tk.LEFT).pack(anchor=tk.W, pady=(2,5))
        self.entry_aas = tk.Entry(fa, font=('Courier New',10), width=24,
                                   state=tk.DISABLED, relief=tk.SOLID, bd=1, bg='#f8fafc')
        self.entry_aas.pack(anchor=tk.W, ipady=3, pady=(0,5))
        fr_anal_btns = tk.Frame(fa, bg=C['green_bg']); fr_anal_btns.pack(anchor=tk.W)
        self.btn_anal = self._btn(fr_anal_btns,'🔍  Analisar',self._analisar_codons,
                                  C['green'],size=9,padx=10,pady=4,state=tk.DISABLED)
        self.btn_auto_codons = self._btn(fr_anal_btns,'⚡  Auto (Preprotein)',
                                         self._analisar_auto_codons,
                                         '#7c3aed',size=9,padx=10,pady=4,state=tk.DISABLED)
        self.lbl_res_aas = tk.Label(fa, text='', font=('Courier New',8),
                                     bg=C['green_bg'], fg=C['texto'],
                                     justify=tk.LEFT, wraplength=310)
        self.lbl_res_aas.pack(anchor=tk.W, pady=(5,0))

        # ── Opção 5' UTR ──
        tk.Frame(fa, bg='#86efac', height=1).pack(fill=tk.X, pady=(8,4))
        fr_utr = tk.Frame(fa, bg=C['green_bg']); fr_utr.pack(anchor=tk.W)
        self.utr_ativo = tk.BooleanVar(value=False)
        tk.Checkbutton(fr_utr, text="Isolar 5' UTR (éxon 1 antes do ATG)",
                       variable=self.utr_ativo, bg=C['green_bg'],
                       font=('Arial',8,'bold'), fg=C['green'],
                       activebackground=C['green_bg'],
                       command=self._toggle_utr).pack(side=tk.LEFT)
        fr_utr2 = tk.Frame(fa, bg=C['green_bg']); fr_utr2.pack(anchor=tk.W, pady=(4,0))
        tk.Label(fr_utr2, text='Cor da 5\'UTR:', font=('Arial',8),
                 bg=C['green_bg'], fg=C['sub']).pack(side=tk.LEFT)
        self.cor_utr = '#6699cc'
        self.btn_cor_utr = tk.Button(fr_utr2, text='  ██  ',
                                      bg='#6699cc', fg='#6699cc',
                                      command=self._escolher_cor_utr,
                                      relief=tk.RAISED, padx=3,
                                      state=tk.DISABLED, cursor='hand2')
        self.btn_cor_utr.pack(side=tk.LEFT, padx=(4,6))
        tk.Label(fr_utr2, text='(azul claro — padrão)',
                 font=('Arial',7,'italic'), bg=C['green_bg'], fg=C['cinza']).pack(side=tk.LEFT)
        self.lbl_utr_info = tk.Label(fa, text='', font=('Arial',7,'italic'),
                                      bg=C['green_bg'], fg=C['sub'])
        self.lbl_utr_info.pack(anchor=tk.W, pady=(2,0))

        # Painel B
        fb = tk.Frame(fr_cod, bg=C['azul_bg'],
                      highlightbackground='#93c5fd', highlightthickness=1,
                      padx=10, pady=8)
        fb.grid(row=0, column=1, sticky='nsew')
        fr_b_tit = tk.Frame(fb, bg=C['azul_bg']); fr_b_tit.pack(fill=tk.X)
        tk.Label(fr_b_tit, text='📍  Localizar códon por nº do aminoácido',
                 font=('Arial',9,'bold'), bg=C['azul_bg'], fg=C['azul']).pack(side=tk.LEFT)
        self._btn_help(fr_b_tit,
            'Encontra o códon de um aminoácido específico\n'
            'pela sua posição na proteína nascente.\n\n'
            'Ex: aminoácido nº 25 = primeiro aa da proteína madura.\n\n'
            'Grifar no Word: marca o códon em ROXO e negrito\n'
            'no documento gerado.',
            bg=C['azul_bg']).pack(side=tk.RIGHT)
        tk.Label(fb, text='Posição na proteína nascente (CDS = éxons concatenados)',
                 font=('Arial',8), bg=C['azul_bg'], fg=C['sub'],
                 justify=tk.LEFT).pack(anchor=tk.W, pady=(2,5))
        fr_ni = tk.Frame(fb, bg=C['azul_bg']); fr_ni.pack(anchor=tk.W)
        tk.Label(fr_ni, text='Aminoácido nº:', bg=C['azul_bg'],
                 font=('Arial',9), fg=C['sub']).pack(side=tk.LEFT)
        self.entry_num = tk.Entry(fr_ni, font=('Courier New',10), width=7,
                                   state=tk.DISABLED, relief=tk.SOLID, bd=1, bg='#f8fafc')
        self.entry_num.pack(side=tk.LEFT, padx=(4,8), ipady=3)
        self.btn_loc = self._btn(fr_ni,'📍  Localizar',self._localizar,
                                 C['azul'],size=9,padx=10,pady=3,state=tk.DISABLED)
        # Opções de grifo no Word
        fr_grifo = tk.Frame(fb, bg=C['azul_bg']); fr_grifo.pack(anchor=tk.W, pady=(6,0))
        tk.Label(fr_grifo, text='Grifar no Word:', bg=C['azul_bg'],
                 font=('Arial',8,'bold'), fg=C['azul']).pack(side=tk.LEFT)
        self.grifo_ativo = tk.BooleanVar(value=True)
        tk.Checkbutton(fr_grifo, text='Sim', variable=self.grifo_ativo,
                       bg=C['azul_bg'], font=('Arial',8),
                       activebackground=C['azul_bg']).pack(side=tk.LEFT, padx=(4,8))
        tk.Label(fr_grifo, text='Cor:', bg=C['azul_bg'],
                 font=('Arial',8), fg=C['sub']).pack(side=tk.LEFT)
        self.cor_grifo = '#7c3aed'   # roxo padrão
        self.btn_cor_grifo = tk.Button(fr_grifo, text='  ██  ',
                                        bg='#7c3aed', fg='#7c3aed',
                                        command=self._escolher_cor_grifo,
                                        relief=tk.RAISED, padx=3, cursor='hand2')
        self.btn_cor_grifo.pack(side=tk.LEFT, padx=(0,4))
        tk.Label(fr_grifo, text='(roxo)', font=('Arial',7,'italic'),
                 bg=C['azul_bg'], fg=C['cinza']).pack(side=tk.LEFT)
        self.lbl_res_cod = tk.Label(fb, text='', font=('Courier New',8),
                                     bg=C['azul_bg'], fg=C['texto'],
                                     justify=tk.LEFT, wraplength=310)
        self.lbl_res_cod.pack(anchor=tk.W, pady=(5,0))

        # ── Painel C: Códon de Parada ──
        STOP_BG  = '#fff1f2'
        STOP_BRD = '#fca5a5'
        STOP_FG  = '#b91c1c'
        fc = tk.Frame(fr_cod, bg=STOP_BG,
                      highlightbackground=STOP_BRD, highlightthickness=1,
                      padx=10, pady=8)
        fc.grid(row=0, column=2, sticky='nsew', padx=(6,0))
        fr_c_tit = tk.Frame(fc, bg='#fff1f2'); fr_c_tit.pack(fill=tk.X)
        tk.Label(fr_c_tit, text='🛑  Códon de Parada',
                 font=('Arial',9,'bold'), bg='#fff1f2', fg=STOP_FG).pack(side=tk.LEFT)
        self._btn_help(fr_c_tit,
            'Localiza o códon de parada (TAA, TAG ou TGA)\n'
            'imediatamente após o último aminoácido da preprotein.\n\n'
            '⚡ Auto (Preprotein): calcula automaticamente usando\n'
            'a sequência da proteína nascente importada do GenBank.\n\n'
            'Grifar no Word: marca o stop codon em VERMELHO e negrito\n'
            'no documento gerado.',
            bg='#fff1f2').pack(side=tk.RIGHT)
        tk.Label(fc, text='Localiza TAA / TAG / TGA após a preprotein',
                 font=('Arial',8), bg=STOP_BG, fg=C['sub'],
                 justify=tk.LEFT).pack(anchor=tk.W, pady=(2,6))
        self.btn_stop = self._btn(fc, '⚡  Auto (Preprotein)',
                                   self._localizar_stop,
                                   STOP_FG, size=9, padx=10, pady=4,
                                   state=tk.DISABLED, anchor=tk.W)
        # Grifo do stop codon
        fr_stop_grifo = tk.Frame(fc, bg=STOP_BG); fr_stop_grifo.pack(anchor=tk.W, pady=(6,0))
        tk.Label(fr_stop_grifo, text='Grifar no Word:', bg=STOP_BG,
                 font=('Arial',8,'bold'), fg=STOP_FG).pack(side=tk.LEFT)
        self.stop_grifo_ativo = tk.BooleanVar(value=True)
        tk.Checkbutton(fr_stop_grifo, text='Sim', variable=self.stop_grifo_ativo,
                       bg=STOP_BG, font=('Arial',8),
                       activebackground=STOP_BG).pack(side=tk.LEFT, padx=(4,8))
        tk.Label(fr_stop_grifo, text='Cor:', bg=STOP_BG,
                 font=('Arial',8), fg=C['sub']).pack(side=tk.LEFT)
        self.cor_stop = '#cc0000'
        self.btn_cor_stop = tk.Button(fr_stop_grifo, text='  ██  ',
                                       bg='#cc0000', fg='#cc0000',
                                       command=self._escolher_cor_stop,
                                       relief=tk.RAISED, padx=3, cursor='hand2')
        self.btn_cor_stop.pack(side=tk.LEFT, padx=(0,4))
        tk.Label(fr_stop_grifo, text='(vermelho)', font=('Arial',7,'italic'),
                 bg=STOP_BG, fg=C['cinza']).pack(side=tk.LEFT)
        self.lbl_res_stop = tk.Label(fc, text='', font=('Courier New',8),
                                      bg=STOP_BG, fg=C['texto'],
                                      justify=tk.LEFT, wraplength=280)
        self.lbl_res_stop.pack(anchor=tk.W, pady=(5,0))

    # ── Botões de geração ────────────────────────

    def _build_btns_gerar(self, parent):
        tk.Frame(parent, bg=C['borda'], height=1).pack(fill=tk.X, pady=(12,0))
        fr = tk.Frame(parent, bg=C['fundo']); fr.pack(pady=10)
        self.btn_gerar = tk.Button(fr,
            text='⚙  Gerar Documento Completo  (Éxons + Íntrons)',
            command=self._gerar_completo, bg=C['cinza'], fg='white',
            font=('Arial',10,'bold'), relief=tk.FLAT, padx=16, pady=8,
            state=tk.DISABLED, cursor='hand2')
        self.btn_gerar.pack(side=tk.LEFT, padx=(0,8))
        self.btn_gerar_ex = tk.Button(fr,
            text='🧬  Gerar Somente Éxons',
            command=self._gerar_so_exons, bg=C['cinza'], fg='white',
            font=('Arial',10,'bold'), relief=tk.FLAT, padx=16, pady=8,
            state=tk.DISABLED, cursor='hand2')
        self.btn_gerar_ex.pack(side=tk.LEFT)
        self.status = tk.StringVar(value='Pronto.')
        tk.Label(parent, textvariable=self.status,
                 font=('Arial',8,'italic'), bg=C['fundo'], fg=C['sub']).pack(pady=(4,12))

    # ── Ações éxons ──────────────────────────────

    def _preencher_proteinas(self, resultado):
        """
        Preenche proteína nascente e madura.
        Nascente = /translation da CDS (já extraída e limpa).
        Madura   = calculada a partir das coordenadas do mat_peptide
                   vs as coordenadas da CDS — para encontrar o offset em aa.
        """
        nasc = resultado.get('proteina_nascente', '')

        # ── Proteína Nascente ──
        self.txt_prot_nasc.config(state=tk.NORMAL)
        self.txt_prot_nasc.delete('1.0', tk.END)
        self.txt_prot_nasc.insert(tk.END, nasc if nasc else '(não encontrada neste registro)')
        self.txt_prot_nasc.config(state=tk.DISABLED)

        # ── Proteína Madura ──
        # A estratégia correta: comparar o início do mat_peptide com o início da CDS
        # para calcular quantos nucleotídeos de offset há → offset_nt / 3 = offset_aa
        mats     = resultado.get('proteina_madura', [])
        proteinas = resultado.get('proteinas', [])
        mad_aa   = ''

        if mats and proteinas and nasc:
            mat_coords = mats[0].get('coords', [])
            cds_blocos = proteinas[0].get('coords_cds', [])  # adicionaremos isso

            # Calcula a posição de início do mat_peptide dentro da CDS (em nt)
            if mat_coords and cds_blocos:
                # Soma dos nt da CDS antes do início do mat_peptide
                mat_inicio_genomico = mat_coords[0][0]
                offset_nt = 0
                for ini_cds, fim_cds in cds_blocos:
                    if mat_inicio_genomico > fim_cds:
                        offset_nt += (fim_cds - ini_cds + 1)
                    elif mat_inicio_genomico >= ini_cds:
                        offset_nt += (mat_inicio_genomico - ini_cds)
                        break
                offset_aa = offset_nt // 3
                mad_aa = nasc[offset_aa:] if offset_aa < len(nasc) else ''
            else:
                # Fallback genérico: remove sinal peptídeo típico (18 aa) + propeptídeo (6 aa)
                # Válido para albumina e proteínas similares; pode variar
                offset_fallback = 24
                mad_aa = nasc[offset_fallback:] if len(nasc) > offset_fallback else nasc

        elif nasc:
            # Sem mat_peptide: usa nascente
            mad_aa = nasc

        self.txt_prot_mad.config(state=tk.NORMAL)
        self.txt_prot_mad.delete('1.0', tk.END)
        self.txt_prot_mad.insert(tk.END, mad_aa if mad_aa else '(não encontrada neste registro)')
        self.txt_prot_mad.config(state=tk.DISABLED)

        n_nasc = len(nasc); n_mad = len(mad_aa)
        self.lbl_prot_info.config(
            text=f'Nascente: {n_nasc} aa   |   Madura: {n_mad} aa', fg=C['green'])

    def _habilitar_exons(self):
        for w in [self.entry_inicio, self.entry_fim, self.spin_tam,
                  self.btn_cor_exon, self.btn_add, self.btn_rem, self.btn_clear,
                  self.btn_ncbi_txt, self.entry_aas, self.entry_num,
                  self.btn_anal, self.btn_loc, self.btn_auto_codons, self.btn_stop]:
            w.config(state=tk.NORMAL)
        self.combo_fonte_exon.config(state='readonly')
        self.btn_add.config(bg='#16a34a')
        self.btn_gerar.config(state=tk.NORMAL, bg=C['header'])
        self.btn_gerar_ex.config(state=tk.NORMAL, bg=C['roxo'])
        self.lbl_bloqueio.config(text='✅  Sequência confirmada. Adicione os Éxons.',fg=C['green'])
        self.lbl_bl_cod.config(text='✅  Pronto para análise de códons.', fg=C['green'])

    def _desabilitar_exons(self):
        for w in [self.entry_inicio, self.entry_fim, self.spin_tam,
                  self.btn_cor_exon, self.btn_add, self.btn_rem, self.btn_clear,
                  self.btn_ncbi_txt, self.entry_aas, self.entry_num,
                  self.btn_anal, self.btn_loc, self.btn_auto_codons, self.btn_stop]:
            w.config(state=tk.DISABLED)
        self.combo_fonte_exon.config(state=tk.DISABLED)
        self.btn_add.config(bg=C['cinza'])
        self.btn_gerar.config(state=tk.DISABLED, bg=C['cinza'])
        self.btn_gerar_ex.config(state=tk.DISABLED, bg=C['cinza'])
        self.lbl_bloqueio.config(text='⚠  Confirme a sequência para habilitar.',fg=C['amber'])
        self.lbl_bl_cod.config(text='⚠  Confirme a sequência e adicione Éxons.',fg=C['amber'])
        self._limpar_exons(confirmar=False)

    def _abrir_importador_ncbi_txt(self):
        JanelaNcbiTxt(self._janela_ref(), self._receber_exons_txt)

    def _receber_exons_txt(self, exons_parseados):
        total = len(self.sequencia_verificada) if self.sequencia_verificada else 0
        for e in exons_parseados:
            if e['fim'] > total: continue
            exon = {'inicio':e['inicio'],'fim':e['fim'],'fonte':self.fonte_exon.get(),
                    'tamanho':self.tamanho_exon.get(),'cor':self.cor_exon_atual,'numero':e['numero']}
            self.exons.append(exon)
            self.tabela.insert('', tk.END, values=(
                f"Éxon {e['numero']}", e['inicio'], e['fim'],
                f"{e['fim']-e['inicio']+1} pb", exon['fonte'], exon['cor']))
        self._atualizar()

    def _adicionar_exon(self):
        try: ini = int(self.entry_inicio.get()); fim = int(self.entry_fim.get())
        except ValueError: messagebox.showerror('Erro','Posições inteiras.'); return
        total = len(self.sequencia_verificada)
        if ini < 1: messagebox.showerror('Erro','Início ≥ 1.'); return
        if ini >= fim: messagebox.showerror('Erro','Início < Fim.'); return
        if fim > total: messagebox.showerror('Erro',f'Fim excede {total:,} pb.'); return
        n = len(self.exons)+1
        exon = {'inicio':ini,'fim':fim,'fonte':self.fonte_exon.get(),
                'tamanho':self.tamanho_exon.get(),'cor':self.cor_exon_atual,'numero':n}
        self.exons.append(exon)
        self.tabela.insert('', tk.END, values=(
            f"Éxon {n}", ini, fim, f"{fim-ini+1} pb", exon['fonte'], exon['cor']))
        self.entry_inicio.delete(0,tk.END); self.entry_fim.delete(0,tk.END)
        self._atualizar()

    def _remover_exon(self):
        sel = self.tabela.selection()
        if not sel: messagebox.showinfo('Atenção','Selecione um éxon.'); return
        idx = self.tabela.index(sel[0])
        self.tabela.delete(sel[0]); self.exons.pop(idx)
        for i, item in enumerate(self.tabela.get_children()):
            v = list(self.tabela.item(item,'values')); v[0]=f"Éxon {i+1}"
            self.tabela.item(item, values=v)
        self._atualizar()

    def _limpar_exons(self, confirmar=True):
        if confirmar and self.exons:
            if not messagebox.askyesno('Confirmar','Remover todos?'): return
        self.exons.clear()
        for i in self.tabela.get_children(): self.tabela.delete(i)
        self._atualizar()

    def _atualizar(self):
        n = len(self.exons)
        self.lbl_cont.config(text=f'{n} éxon(s)' if n > 0 else '')
        self._draw()

    # ── Análise de códons ────────────────────────

    def _toggle_utr(self):
        estado = tk.NORMAL if self.utr_ativo.get() else tk.DISABLED
        self.btn_cor_utr.config(state=estado)
        if not self.utr_ativo.get():
            self.lbl_utr_info.config(text='')
            self._utr_regiao = None

    def _escolher_cor_utr(self):
        cor = colorchooser.askcolor(color=self.cor_utr, title="Cor da 5' UTR")[1]
        if cor:
            self.cor_utr = cor
            self.btn_cor_utr.configure(bg=cor, fg=cor)

    def _escolher_cor_stop(self):
        cor = colorchooser.askcolor(color=self.cor_stop, title='Cor do códon de parada')[1]
        if cor:
            self.cor_stop = cor
            self.btn_cor_stop.configure(bg=cor, fg=cor)

    def _localizar_stop(self):
        """Localiza o códon de parada usando os últimos aa da proteína nascente."""
        if not self.sequencia_verificada or not self.exons:
            messagebox.showerror('Erro', 'Adicione os Éxons primeiro.'); return
        nasc = ''
        if hasattr(self, 'txt_prot_nasc'):
            nasc = self.txt_prot_nasc.get('1.0', tk.END).strip()
            if nasc.startswith('('): nasc = ''
        if not nasc:
            messagebox.showinfo('Info',
                'Proteína nascente não disponível.\n'
                'Importe um gene do GenBank primeiro.'); return
        coords_cds = None
        offset     = getattr(self, '_offset_regiao', 0)
        if hasattr(self, '_proteinas_genbank') and self._proteinas_genbank:
            coords_cds = self._proteinas_genbank[0].get('coords_cds')
        res = localizar_stop_codon(
            self.sequencia_verificada, self.exons,
            proteina_nascente=nasc, coords_cds=coords_cds, offset=offset)
        if res.get('erro'):
            self.lbl_res_stop.config(text=f"❌ {res['erro']}", fg='#dc2626'); return
        if not res['encontrado']:
            aviso = res.get('aviso', '')
            self.lbl_res_stop.config(
                text=f"⚠️  {aviso}\nCódon: {res['codon']} "
                     f"(pos. {res['pos_inicio']}–{res['pos_fim']})",
                fg=C['amber']); return
        ultimos = ' – '.join(res['ultimos_aas']) if res['ultimos_aas'] else ''
        texto = (f"✅  Stop codon: {res['codon']}\n"
                 f"   Posição: {res['pos_inicio']:,}–{res['pos_fim']:,} pb\n"
                 f"   Após aa nº {res['n_aa_total']} ({ultimos})")
        self.lbl_res_stop.config(text=texto, fg='#b91c1c')
        # Marcador stop na barra (vermelho)
        self.marc_stop = {
            'pi':    res['pos_inicio'] - 1,
            'pf':    res['pos_fim'],
            'label': f"Stop ({res['codon']})",
            'cor':   '#dc2626',
        }
        # Grifo no Word
        if self.stop_grifo_ativo.get():
            self._stop_grifo = {
                'pos_ini':   res['pos_inicio'] - 1,  # 0-based
                'pos_fim':   res['pos_fim'],           # exclusivo
                'cor':       self.cor_stop,
                'negrito':   True,
            }
        else:
            self._stop_grifo = None
        self._draw()

    def _escolher_cor_grifo(self):
        cor = colorchooser.askcolor(color=self.cor_grifo, title='Cor do grifo')[1]
        if cor:
            self.cor_grifo = cor
            self.btn_cor_grifo.configure(bg=cor, fg=cor)

    def _analisar_auto_codons(self):
        """
        Pega automaticamente os primeiros N aminoácidos da proteína nascente
        e usa como entrada para analisar os primeiros códons na CDS.
        """
        if not hasattr(self, 'txt_prot_nasc'):
            messagebox.showinfo('Info', 'Importe um gene do GenBank primeiro para ter a proteína nascente.')
            return
        nasc = self.txt_prot_nasc.get('1.0', tk.END).strip()
        if not nasc or nasc.startswith('('):
            messagebox.showinfo('Info', 'Proteína nascente não disponível. Importe do GenBank primeiro.')
            return
        if not self.exons:
            messagebox.showerror('Erro', 'Adicione os Éxons primeiro.')
            return
        # Usa os primeiros 6 aminoácidos para análise robusta
        n_usar = min(6, len(nasc))
        aas = [AA_SINONIMOS.get(c.upper(), c.upper()) for c in nasc[:n_usar]
               if c.upper() in AA_SINONIMOS]
        if len(aas) < 3:
            messagebox.showerror('Erro', 'Proteína nascente muito curta para análise.')
            return
        # Preenche o campo e executa
        self.entry_aas.delete(0, tk.END)
        self.entry_aas.insert(0, ','.join(aas))
        self._analisar_codons()

    def _analisar_codons(self):
        if not self.sequencia_verificada or not self.exons:
            messagebox.showerror('Erro','Adicione os Éxons primeiro.'); return
        entrada = self.entry_aas.get().strip()
        if not entrada: messagebox.showerror('Erro','Informe os aminoácidos.'); return
        partes = [p for p in re.split(r'[,\s]+', entrada) if p.strip()]
        if len(partes) < 3: messagebox.showerror('Erro','Mínimo 3 aminoácidos.'); return
        # Obtém proteína nascente e coords_cds se disponíveis (modo NCBI)
        prot_nasc  = ''
        coords_cds = None
        offset     = getattr(self, '_offset_regiao', 0)
        if hasattr(self, 'txt_prot_nasc'):
            prot_nasc = self.txt_prot_nasc.get('1.0', tk.END).strip()
            if prot_nasc.startswith('('): prot_nasc = ''
        if hasattr(self, '_proteinas_genbank') and self._proteinas_genbank:
            coords_cds = self._proteinas_genbank[0].get('coords_cds')
        res = analisar_codons(self.sequencia_verificada, self.exons, partes,
                              proteina_nascente=prot_nasc,
                              coords_cds=coords_cds, offset=offset)
        if res.get('erro'):
            self.lbl_res_aas.config(text=f"❌ {res['erro']}", fg='#dc2626'); return
        matches = res.get('matches', [True]*len(partes))
        n_ok  = sum(matches)
        n_tot = len(matches)
        fonte = res.get('fonte','exons')
        fonte_txt = ' (via /translation GenBank)' if fonte == 'translation' else ' (via éxons)'
        if res['encontrado']:
            linhas = [f'✅ Todos os {n_tot} aminoácidos confirmados{fonte_txt}\n']
            for i,(aa,cod,pos) in enumerate(zip(
                    res['aas_fornecidos'],res['combinacao'],res['posicoes_gene'])):
                pos_txt = f'pos. {pos} no gene' if pos > 0 else 'pos. calculada'
                linhas.append(f'  Códon {i+1}: {cod}  →  {aa}  ({pos_txt})')
            self.lbl_res_aas.config(text='\n'.join(linhas), fg=C['green'])
            self.marc_inicio = None  # painel A não gera marcador na barra
            # ── Calcular 5' UTR ──
            # UTR = nucleotídeos do éxon 1 antes do ATG
            self._utr_regiao = None
            if self.exons and res['pos_gene_inicio'] > 0:
                exon1_ini  = self.exons[0]['inicio']      # 1-based
                atg_pos    = res['pos_gene_inicio']        # 1-based, posição do A do ATG
                if atg_pos > exon1_ini:
                    utr_len = atg_pos - exon1_ini
                    self._utr_regiao = {
                        'pos_ini': exon1_ini - 1,   # 0-based
                        'pos_fim': atg_pos - 1,      # 0-based exclusivo
                        'cor':     self.cor_utr,
                        'len':     utr_len,
                    }
                    if hasattr(self, 'utr_ativo') and self.utr_ativo.get():
                        self.lbl_utr_info.config(
                            text=f"5' UTR detectada: {utr_len} nt "
                                 f"(pos. {exon1_ini}–{atg_pos-1})",
                            fg=C['green'])
                else:
                    if hasattr(self, 'lbl_utr_info'):
                        self.lbl_utr_info.config(
                            text="ATG no início do éxon 1 — sem 5' UTR.", fg=C['sub'])
        else:
            linhas = [f'⚠  {n_ok}/{n_tot} aminoácidos correspondem.\n','Detalhes:']
            for i,(af,ac,cod,ok) in enumerate(zip(
                    res['aas_fornecidos'],res['aas_na_cds'],
                    res['codons_na_cds'],matches)):
                m = '✅' if ok else '❌'
                linhas.append(f'  {m} Pos {i+1}: {cod} → {ac}  (fornecido: {af})')
            if n_ok == 0:
                linhas.append('\n💡 Verifique se os éxons estão corretos ou se a região carregada inclui o início do gene.')
            self.lbl_res_aas.config(text='\n'.join(linhas), fg=C['amber'])
        self._draw()

    def _localizar(self):
        if not self.sequencia_verificada or not self.exons:
            messagebox.showerror('Erro','Adicione os Éxons primeiro.'); return
        try: num = int(self.entry_num.get().strip())
        except ValueError: messagebox.showerror('Erro','Número inteiro.'); return
        res = localizar_codon(self.sequencia_verificada, self.exons, num)
        if res.get('erro'):
            self.lbl_res_cod.config(text=f"❌ {res['erro']}", fg='#dc2626'); return
        pos_i = res['pos_gene'][0]; pos_f = res['pos_gene'][2]
        grifo_info = ''
        if hasattr(self,'grifo_ativo') and self.grifo_ativo.get():
            grifo_info = f'  |  🖊 Grifado em {self.cor_grifo} no próximo documento'
        texto = (f"✅  Aminoácido {num}: {res['aminoacido']}  ({res['codon']})\n"
                 f"   Posição no gene: {res['pos_inicio']}–{res['pos_fim']} pb{grifo_info}")
        self.lbl_res_cod.config(text=texto, fg=C['azul'])
        self.marc_loc = {'pi':   pos_i, 'pf': pos_f,
                         'label': f"aa {num} ({res['aminoacido']})",
                         'cor':   '#7c3aed'}  # roxo
        # Salva grifo para usar na geração do documento
        if hasattr(self,'grifo_ativo') and self.grifo_ativo.get():
            self._codon_grifo = {
                'pos_ini':  pos_i - 1,   # 0-based
                'pos_fim':  pos_f,        # exclusivo
                'cor':      self.cor_grifo,
                'negrito':  True,
                'num_aa':   num,
                'aminoacido': res['aminoacido'],
                'codon':    res['codon'],
            }
        else:
            self._codon_grifo = None
        self._draw()

    # ── Gerar documentos ─────────────────────────

    def _gerar_completo(self):
        if not self.sequencia_verificada: messagebox.showerror('Erro','Confirme a sequência.'); return
        path = filedialog.asksaveasfilename(defaultextension='.docx',
            filetypes=[('Word','*.docx')], title='Salvar documento completo…')
        if not path: return
        self.status.set('Processando…'); self._janela_ref().update()
        try:
            cfg = {'fonte':self.fonte_base.get(),'cor':self.cor_base,'tamanho':self.tamanho_base.get()}
            codons_g = []
            if getattr(self,'_codon_grifo',None): codons_g.append(self._codon_grifo)
            if getattr(self,'_stop_grifo',None):  codons_g.append(self._stop_grifo)
            utr_g = None
            if (hasattr(self,'utr_ativo') and self.utr_ativo.get()
                    and getattr(self,'_utr_regiao',None)):
                utr_g = dict(self._utr_regiao)
                utr_g['cor'] = self.cor_utr
            n = gerar_completo(self.sequencia_verificada, cfg, self.exons, path,
                               caixa=self.caixa.get(), codons_grifar=codons_g, utr_regiao=utr_g)
            self.status.set(f'✅ {n:,} pb, {len(self.exons)} éxon(s).')
            messagebox.showinfo('Sucesso',f'Documento salvo!\n\n• {n:,} pb\n• {len(self.exons)} éxon(s)\n\n{path}')
        except Exception as e: messagebox.showerror('Erro',str(e))

    def _gerar_so_exons(self):
        if not self.sequencia_verificada: messagebox.showerror('Erro','Confirme a sequência.'); return
        if not self.exons: messagebox.showerror('Erro','Adicione ao menos um Éxon.'); return
        path = filedialog.asksaveasfilename(defaultextension='.docx',
            filetypes=[('Word','*.docx')], title='Salvar somente éxons…')
        if not path: return
        self.status.set('Processando…'); self._janela_ref().update()
        try:
            # Monta codons_grifar e utr_regiao para compatibilidade
            codons_g = []
            if getattr(self,'_codon_grifo',None): codons_g.append(self._codon_grifo)
            if getattr(self,'_stop_grifo',None):  codons_g.append(self._stop_grifo)
            utr_g = None
            if (hasattr(self,'utr_ativo') and self.utr_ativo.get()
                    and getattr(self,'_utr_regiao',None)):
                utr_g = dict(self._utr_regiao)
                utr_g['cor'] = self.cor_utr
            n = gerar_so_exons(self.sequencia_verificada, self.exons, path,
                               caixa=self.caixa.get(),
                               codons_grifar=codons_g, utr_regiao=utr_g)
            self.status.set(f'✅ Somente éxons — {n:,} pb.')
            messagebox.showinfo('Sucesso',
                f'Documento salvo!\n\n• {n:,} pb codificantes\n• Éxons ímpares: preto\n• Éxons pares: laranja\n\n{path}')
        except Exception as e: messagebox.showerror('Erro',str(e))


# ══════════════════════════════════════════════════════════════
# JANELA NCBI TEXTO (importador de texto colado)
# ══════════════════════════════════════════════════════════════

class JanelaNcbiTxt(tk.Toplevel):
    def __init__(self, parent, callback):
        super().__init__(parent)
        self.title('📋 Importar texto do NCBI'); self.geometry('680x500')
        self.configure(bg=C['fundo']); self.grab_set()
        self.callback = callback; self.exons_ok = []
        self._build()

    def _build(self):
        tk.Label(self, text='📋  Importar Éxons — Texto do GenBank',
                 font=('Arial',12,'bold'), bg=C['roxo'], fg='white', pady=10).pack(fill=tk.X)
        tk.Label(self, text="Cole o texto da seção 'Features' do NCBI GenBank abaixo.",
                 font=('Arial',9), bg=C['fundo'], fg=C['sub'], pady=6).pack()
        fr = tk.Frame(self, bg=C['fundo'], padx=14); fr.pack(fill=tk.BOTH, expand=True)
        self.txt = tk.Text(fr, font=('Courier New',9), height=12,
                           wrap=tk.WORD, relief=tk.SOLID, bd=1, bg='white')
        sc = ttk.Scrollbar(fr, command=self.txt.yview)
        self.txt.configure(yscrollcommand=sc.set)
        self.txt.pack(side=tk.LEFT, fill=tk.BOTH, expand=True); sc.pack(side=tk.RIGHT, fill=tk.Y)
        tk.Button(self, text='🔍  Identificar Éxons', command=self._parsear,
                  bg=C['roxo'], fg='white', font=('Arial',10,'bold'),
                  relief=tk.FLAT, padx=14, pady=6).pack(pady=6)
        self.lbl = tk.Label(self, text='', font=('Arial',9,'italic'),
                             bg=C['fundo'], fg=C['sub']); self.lbl.pack()
        fr2 = tk.Frame(self, bg=C['fundo'], padx=14); fr2.pack(fill=tk.X)
        cols = ('Éxon','Início','Fim','Tamanho')
        self.tab = ttk.Treeview(fr2, columns=cols, show='headings', height=4)
        for col in cols:
            self.tab.heading(col,text=col); self.tab.column(col,width=155,anchor=tk.CENTER)
        self.tab.pack(fill=tk.X)
        fr3 = tk.Frame(self, bg=C['fundo']); fr3.pack(pady=8)
        self.btn_imp = tk.Button(fr3, text='✅  Importar',
                                  command=self._importar, bg=C['green'], fg='white',
                                  font=('Arial',10,'bold'), relief=tk.FLAT,
                                  padx=14, pady=6, state=tk.DISABLED)
        self.btn_imp.pack(side=tk.LEFT, padx=(0,8))
        tk.Button(fr3, text='Cancelar', command=self.destroy,
                  bg=C['cinza'], fg='white', font=('Arial',9),
                  relief=tk.FLAT, padx=12, pady=6).pack(side=tk.LEFT)

    def _parsear(self):
        for i in self.tab.get_children(): self.tab.delete(i)
        exons = parsear_ncbi_texto(self.txt.get('1.0', tk.END))
        if not exons:
            self.lbl.config(text='⚠  Nenhum éxon encontrado.', fg=C['amber'])
            self.btn_imp.config(state=tk.DISABLED); return
        for e in exons:
            self.tab.insert('', tk.END, values=(
                f"Éxon {e['numero']}", e['inicio'], e['fim'], f"{e['fim']-e['inicio']+1} pb"))
        self.exons_ok = exons
        self.lbl.config(text=f"✅  {len(exons)} éxon(s) identificado(s).", fg=C['green'])
        self.btn_imp.config(state=tk.NORMAL)

    def _importar(self):
        if self.exons_ok: self.callback(self.exons_ok); self.destroy()


# ══════════════════════════════════════════════════════════════
# MODO AUTOMÁTICO — NCBI API
# ══════════════════════════════════════════════════════════════

class ModoAutomatico(tk.Toplevel, EditorMixin):
    def __init__(self, parent):
        super().__init__(parent)
        self.title('🧬 ExonEditor v7.3 — Modo Automático (NCBI API)')
        self.geometry('980x860'); self.minsize(860,660)
        self.configure(bg=C['fundo'])

        # Estado
        self.sequencia_verificada = None
        self.exons                = []
        self.cor_base             = '#aaaaaa'
        self.cor_exon_atual       = '#000000'
        self.fonte_base           = tk.StringVar(value='Courier New')
        self.tamanho_base         = tk.IntVar(value=11)
        self.fonte_exon           = tk.StringVar(value='Courier New')
        self.tamanho_exon         = tk.IntVar(value=11)
        self.caixa                = tk.StringVar(value='minuscula')
        self.email_ncbi           = tk.StringVar(value='')
        self.registros_busca      = []
        self._proteinas_genbank   = []
        self._offset_regiao       = 0
        self._codon_grifo         = None
        self.cor_grifo            = '#7c3aed'
        self._utr_regiao          = None
        self.cor_utr              = '#6699cc'
        self._stop_grifo          = None
        self.cor_stop             = '#cc0000'
        self.marc_inicio          = None
        self.marc_stop            = None
        self.marc_loc             = None
        self._build()

    def _janela_ref(self): return self

    def _build(self):
        # Header
        hdr = tk.Frame(self, bg='#0f2744'); hdr.pack(fill=tk.X)
        tk.Label(hdr, text='🌐  Modo Automático — NCBI API',
                 font=('Arial',14,'bold'), bg='#0f2744', fg='white',
                 pady=10, padx=16).pack(side=tk.LEFT)
        tk.Button(hdr, text='← Voltar ao Início', command=self._voltar,
                  bg='#1e4a7a', fg='white', font=('Arial',9),
                  relief=tk.FLAT, padx=10, pady=6,
                  cursor='hand2').pack(side=tk.RIGHT, padx=12, pady=8)

        # Scroll
        outer = tk.Frame(self, bg=C['fundo']); outer.pack(fill=tk.BOTH, expand=True)
        cv = tk.Canvas(outer, bg=C['fundo'], highlightthickness=0)
        vsb = ttk.Scrollbar(outer, orient=tk.VERTICAL, command=cv.yview)
        cv.configure(yscrollcommand=vsb.set)
        vsb.pack(side=tk.RIGHT, fill=tk.Y); cv.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.inner = tk.Frame(cv, bg=C['fundo'])
        wid = cv.create_window((0,0), window=self.inner, anchor='nw')
        cv.bind('<Configure>', lambda e: cv.itemconfig(wid, width=e.width))
        self.inner.bind('<Configure>', lambda e: cv.configure(scrollregion=cv.bbox('all')))
        def _scroll_auto(e):
            try: cv.yview_scroll(int(-1*(e.delta/120)),'units')
            except Exception: pass
        cv.bind('<MouseWheel>', _scroll_auto)
        self.inner.bind_all('<MouseWheel>', _scroll_auto)
        self.protocol('WM_DELETE_WINDOW', lambda: (self.inner.unbind_all('<MouseWheel>'), self._voltar()))

        p = tk.Frame(self.inner, bg=C['fundo'], padx=18, pady=12)
        p.pack(fill=tk.BOTH, expand=True)

        self._build_busca_ncbi(p)
        self._build_fmt_base(p)
        self._build_tabela_exons(p)
        self._build_barra(p)
        self._build_codons(p)
        self._build_btns_gerar(p)

    def _build_busca_ncbi(self, p):
        c = self._card(p, '① Buscar Gene no NCBI GenBank',
                       cor_t='white', bg='#0f2744', pady_top=0)
        c.master.configure(highlightbackground='#1e4a7a')

        # Email
        fr_e = tk.Frame(c, bg='#0f2744'); fr_e.pack(fill=tk.X, pady=(0,8))
        tk.Label(fr_e, text='Email NCBI:', font=('Arial',9),
                 bg='#0f2744', fg='#93c5fd').pack(side=tk.LEFT)
        self.entry_email = tk.Entry(fr_e, textvariable=self.email_ncbi,
                                    font=('Arial',10), width=30,
                                    relief=tk.SOLID, bd=1, bg='#1e3a5f', fg='white',
                                    insertbackground='white')
        self.entry_email.pack(side=tk.LEFT, padx=(6,12), ipady=4)
        tk.Label(fr_e, text='(recomendado para mais requisições/seg)',
                 font=('Arial',8,'italic'), bg='#0f2744', fg='#64748b').pack(side=tk.LEFT)

        # Busca
        fr_b = tk.Frame(c, bg='#0f2744'); fr_b.pack(fill=tk.X, pady=(0,8))
        tk.Label(fr_b, text='Gene / Accession:', font=('Arial',9),
                 bg='#0f2744', fg='#93c5fd').pack(side=tk.LEFT)
        self.entry_busca = tk.Entry(fr_b, font=('Arial',11), width=22,
                                    relief=tk.SOLID, bd=1, bg='#1e3a5f', fg='white',
                                    insertbackground='white')
        self.entry_busca.pack(side=tk.LEFT, padx=(6,8), ipady=4)
        self.entry_busca.bind('<Return>', lambda e: self._buscar())
        tk.Label(fr_b, text='Organismo:', font=('Arial',9),
                 bg='#0f2744', fg='#93c5fd').pack(side=tk.LEFT)
        self.entry_org = tk.Entry(fr_b, font=('Arial',11), width=16,
                                   relief=tk.SOLID, bd=1, bg='#1e3a5f', fg='white',
                                   insertbackground='white')
        self.entry_org.insert(0, 'Homo sapiens')
        self.entry_org.pack(side=tk.LEFT, padx=(6,8), ipady=4)
        tk.Button(fr_b, text='🔍  Buscar', command=self._buscar,
                  bg='#2563eb', fg='white', font=('Arial',10,'bold'),
                  relief=tk.FLAT, padx=12, pady=5, cursor='hand2').pack(side=tk.LEFT)

        fr_busca_tit = tk.Frame(c, bg='#0f2744'); fr_busca_tit.pack(fill=tk.X, pady=(0,2))
        self.lbl_status_busca = tk.Label(fr_busca_tit, text='Digite o gene e clique em Buscar.',
                                         font=('Arial',8,'italic'),
                                         bg='#0f2744', fg='#64748b')
        self.lbl_status_busca.pack(side=tk.LEFT)
        self._btn_help(fr_busca_tit,
            'Busca genes no NCBI GenBank por nome, símbolo ou accession.\n\n'
            'Gene / Accession: Ex: ALB, TP53, NM_000477, NG_009291\n'
            'Organismo: Homo sapiens (padrão)\n\n'
            'O app busca sempre registros de Gene Genômico Completo\n'
            '(NG_ / NC_) que contêm éxons e íntrons anotados.\n\n'
            'Após a busca, selecione o registro e clique em Carregar.\n'
            'O app detecta a região automaticamente em 2 passos.',
            bg='#0f2744').pack(side=tk.RIGHT)

        # Tabela de resultados
        fr_res = tk.Frame(c, bg='#0f2744'); fr_res.pack(fill=tk.X)
        cols = ('Accession', 'Título', 'Tamanho (pb)')
        self.tab_res = ttk.Treeview(fr_res, columns=cols, show='headings', height=4)
        self.tab_res.heading('Accession',   text='Accession')
        self.tab_res.heading('Título',      text='Título do Registro')
        self.tab_res.heading('Tamanho (pb)',text='Tamanho (pb)')
        self.tab_res.column('Accession',    width=130, anchor=tk.CENTER)
        self.tab_res.column('Título',       width=420)
        self.tab_res.column('Tamanho (pb)', width=110, anchor=tk.CENTER)
        sc = ttk.Scrollbar(fr_res, orient=tk.VERTICAL, command=self.tab_res.yview)
        self.tab_res.configure(yscrollcommand=sc.set)
        self.tab_res.pack(side=tk.LEFT, fill=tk.X, expand=True)
        sc.pack(side=tk.RIGHT, fill=tk.Y)
        # Ao selecionar registro: habilita botão e limpa campos De/Até
        self.tab_res.bind('<<TreeviewSelect>>', self._on_select_registro)

        # Info de busca
        tk.Label(c, text='🧬  Gene Genômico Completo (NG_ / NC_)',
                 font=('Arial',8,'italic'), bg='#0f2744', fg='#93c5fd').pack(
                 anchor=tk.W, pady=(8,0))

        # Campos De / Até (preenchidos automaticamente após 1ª carga)
        fr_reg = tk.Frame(c, bg='#0f2744'); fr_reg.pack(fill=tk.X, pady=(8,0))
        tk.Label(fr_reg, text='Região detectada:',
                 font=('Arial',9), bg='#0f2744', fg='#93c5fd').pack(side=tk.LEFT)
        tk.Label(fr_reg, text='De:', font=('Arial',9),
                 bg='#0f2744', fg='#93c5fd').pack(side=tk.LEFT, padx=(8,0))
        self.entry_from = tk.Entry(fr_reg, font=('Arial',10), width=8,
                                   relief=tk.SOLID, bd=1, bg='#1e3a5f', fg='white',
                                   insertbackground='white', state=tk.DISABLED)
        self.entry_from.pack(side=tk.LEFT, padx=(4,8), ipady=3)
        tk.Label(fr_reg, text='Até:', font=('Arial',9),
                 bg='#0f2744', fg='#93c5fd').pack(side=tk.LEFT)
        self.entry_to = tk.Entry(fr_reg, font=('Arial',10), width=8,
                                  relief=tk.SOLID, bd=1, bg='#1e3a5f', fg='white',
                                  insertbackground='white', state=tk.DISABLED)
        self.entry_to.pack(side=tk.LEFT, padx=(4,0), ipady=3)
        tk.Label(fr_reg, text='pb', font=('Arial',8,'italic'),
                 bg='#0f2744', fg='#64748b').pack(side=tk.LEFT, padx=(4,0))
        self.lbl_prev_pb = tk.Label(fr_reg, text='',
                                     font=('Arial',9,'bold'),
                                     bg='#0f2744', fg='#4ade80')
        self.lbl_prev_pb.pack(side=tk.LEFT, padx=(10,0))
        self.entry_from.bind('<KeyRelease>', lambda e: self._atualizar_preview_pb())
        self.entry_to.bind('<KeyRelease>',   lambda e: self._atualizar_preview_pb())

        # Botão único de carregar
        fr_btn = tk.Frame(c, bg='#0f2744'); fr_btn.pack(anchor=tk.W, pady=(10,0))
        self.btn_carregar = tk.Button(fr_btn,
                                      text='📥  Carregar Registro',
                                      command=self._carregar,
                                      bg='#7c3aed', fg='white',
                                      font=('Arial',10,'bold'),
                                      relief=tk.FLAT, padx=14, pady=6,
                                      state=tk.DISABLED, cursor='hand2')
        self.btn_carregar.pack(side=tk.LEFT)

        self.lbl_importado = tk.Label(c, text='',
                                       font=('Arial',8,'italic'),
                                       bg='#0f2744', fg='#4ade80')
        self.lbl_importado.pack(anchor=tk.W, pady=(4,0))

    # ── Lógica de busca e carregamento ──────────────────────────

    def _on_select_registro(self, event):
        """Ao selecionar um registro na tabela: habilita botão e limpa De/Até."""
        self.btn_carregar.config(state=tk.NORMAL)
        # Limpa campos e estado do gene anterior
        for w in (self.entry_from, self.entry_to):
            w.config(state=tk.NORMAL)
            w.delete(0, tk.END)
            w.config(state=tk.DISABLED)
        self.lbl_prev_pb.config(text='')
        self._acc_atual   = None
        self._email_atual = None

    def _atualizar_preview_pb(self):
        try:
            de  = int(self.entry_from.get().strip())
            ate = int(self.entry_to.get().strip())
            if ate > de:
                self.lbl_prev_pb.config(text=f'→  {ate - de + 1:,} pb nessa região')
            else:
                self.lbl_prev_pb.config(text='')
        except (ValueError, AttributeError):
            self.lbl_prev_pb.config(text='')

    def _buscar(self):
        termo = self.entry_busca.get().strip()
        if not termo:
            messagebox.showerror('Erro', 'Digite um gene ou accession.')
            return
        org   = self.entry_org.get().strip()
        email = self.email_ncbi.get().strip() or 'anonymous@exoneditor.tool'
        if org:
            query = (f'({termo}[Gene Name] OR {termo}[Accession])'
                     f' AND {org}[Organism] AND RefSeqGene[Filter]')
        else:
            query = f'({termo}[Gene Name] OR {termo}[Accession]) AND RefSeqGene[Filter]'
        self.lbl_status_busca.config(text='Buscando… aguarde.', fg='#fbbf24')
        self.btn_carregar.config(state=tk.DISABLED)
        self.update()

        def _t():
            try:
                regs = ncbi_buscar(query, email)
                self.after(0, lambda: self._exibir_resultados(regs))
            except Exception as e:
                msg = str(e)
                self.after(0, lambda: self.lbl_status_busca.config(
                    text=f'Erro: {msg}', fg='#f87171'))
        threading.Thread(target=_t, daemon=True).start()

    def _exibir_resultados(self, regs):
        for i in self.tab_res.get_children():
            self.tab_res.delete(i)
        self.registros_busca = regs
        if not regs:
            self.lbl_status_busca.config(text='Nenhum resultado.', fg='#fbbf24')
            return
        for r in regs:
            tam = f'{int(r["len"]):,}' if str(r["len"]).isdigit() else r["len"]
            self.tab_res.insert('', tk.END,
                                values=(r['acc'], r['titulo'][:75], tam))
        self.lbl_status_busca.config(
            text=f'{len(regs)} resultado(s). Selecione e clique em Carregar.',
            fg='#4ade80')

    def _carregar(self):
        """
        Fluxo de 2 passos (transparente ao usuário):
          Passo 1 — baixa sequência completa → detecta região dos éxons
          Passo 2 — baixa apenas a região → importa com coordenadas corretas
        Se campos De/Até já estiverem preenchidos (passo 2 automático ou manual),
        vai direto para o passo 2.
        """
        sel = self.tab_res.selection()
        if not sel:
            messagebox.showinfo('Atenção', 'Selecione um registro na tabela.')
            return
        idx   = self.tab_res.index(sel[0])
        acc   = self.registros_busca[idx]['acc']
        email = self.email_ncbi.get().strip() or 'anonymous@exoneditor.tool'

        # Guarda para uso no passo 2 (independe da seleção da tabela)
        self._acc_atual   = acc
        self._email_atual = email

        de_txt  = self.entry_from.get().strip()
        ate_txt = self.entry_to.get().strip()

        if de_txt and ate_txt:
            # Campos já preenchidos → passo 2 direto
            try:
                seq_start = int(de_txt)
                seq_stop  = int(ate_txt)
                assert seq_start >= 1 and seq_stop > seq_start
            except (ValueError, AssertionError):
                messagebox.showerror('Erro', 'Valores De/Até inválidos.')
                return
            self._fetch(acc, email, seq_start, seq_stop, passo=2)
        else:
            # Campos vazios → passo 1 (baixa completo para detectar região)
            self._fetch(acc, email, None, None, passo=1)

    def _fetch(self, acc, email, seq_start, seq_stop, passo):
        """Executa download em thread. passo=1 → completo; passo=2 → região."""
        info = f'[{seq_start:,}–{seq_stop:,} pb]' if seq_start else '[sequência completa]'
        self.lbl_status_busca.config(
            text=f'Passo {passo}/2 — Carregando {acc} {info}…', fg='#fbbf24')
        self.btn_carregar.config(state=tk.DISABLED)
        self.update()

        def _t():
            try:
                gb  = ncbi_fetch_genbank(acc, email,
                                          seq_start=seq_start, seq_stop=seq_stop)
                res = parsear_genbank(gb)
                res['_passo'] = passo
                self.after(0, lambda: self._importar_genbank(res))
            except Exception as e:
                msg = str(e)
                self.after(0, lambda: (
                    self.lbl_status_busca.config(
                        text=f'Erro: {msg}', fg='#f87171'),
                    self.btn_carregar.config(state=tk.NORMAL)))
        threading.Thread(target=_t, daemon=True).start()

    def _importar_genbank(self, res):
        """Importa resultado do GenBank."""
        self.btn_carregar.config(state=tk.NORMAL)
        seq   = res.get('sequencia', '')
        passo = res.get('_passo', 2)
        if not seq:
            messagebox.showerror('Erro', 'Sequência vazia.')
            return

        # ── Passo 1: detectou região → preenche campos e dispara passo 2 ──
        if passo == 1:
            self.marc_inicio = None  # limpa marcadores do gene anterior
            self.marc_stop   = None
            self.marc_loc    = None
            reg = res.get('regiao_sugerida')
            if not reg:
                messagebox.showerror('Erro',
                    'Não foi possível detectar a região dos éxons.\n'
                    'Preencha os campos De/Até manualmente e clique em Carregar.')
                return
            de_val, ate_val = reg['de'], reg['ate']
            # Preenche campos
            for w in (self.entry_from, self.entry_to):
                w.config(state=tk.NORMAL)
            self.entry_from.delete(0, tk.END); self.entry_from.insert(0, str(de_val))
            self.entry_to.delete(0, tk.END);   self.entry_to.insert(0, str(ate_val))
            self._atualizar_preview_pb()
            pb = ate_val - de_val + 1
            self.lbl_status_busca.config(
                text=f'Região detectada: {de_val:,}–{ate_val:,} ({pb:,} pb). '
                     f'Carregando região…', fg='#fbbf24')
            # Dispara passo 2 diretamente (sem after, sem depender de seleção)
            acc   = self._acc_atual
            email = self._email_atual
            self._fetch(acc, email, de_val, ate_val, passo=2)
            return

        # ── Passo 2: importa éxons e finaliza ──
        total = len(seq)
        self.sequencia_verificada = None
        self.marc_inicio = None
        self.marc_stop   = None
        self.marc_loc    = None
        self._limpar_exons(confirmar=False)
        self._desabilitar_exons()
        self.sequencia_verificada = seq
        self._habilitar_exons()

        for e in res.get('exons', []):
            ini, fim = e['inicio'], e['fim']
            if ini < 1 or fim > total:
                continue
            exon = {'inicio': ini, 'fim': fim,
                    'fonte':   self.fonte_exon.get(),
                    'tamanho': self.tamanho_exon.get(),
                    'cor':     self.cor_exon_atual,
                    'numero':  e['numero']}
            self.exons.append(exon)
            self.tabela.insert('', tk.END, values=(
                f"Éxon {e['numero']}", ini, fim,
                f"{fim-ini+1} pb", exon['fonte'], exon['cor']))

        self._atualizar()
        self._proteinas_genbank = res.get('proteinas', [])
        self._preencher_proteinas(res)

        acc  = res.get('accession', '')
        n_ex = len(self.exons)
        n_pr = len(res.get('proteinas', []))
        self.lbl_status_busca.config(
            text=f'✅ {acc} — {total:,} pb, {n_ex} éxon(s) importados.',
            fg='#4ade80')
        self.lbl_importado.config(
            text=f'✅  {acc}  |  {total:,} pb  |  {n_ex} éxon(s)  |  {n_pr} proteína(s)')
        prots = res.get('proteinas', [])
        if prots:
            p0 = prots[0]
            self.status.set(
                f"{acc} — {p0['produto']} ({p0['n_aa']} aa) + {n_ex} éxons importados.")
        else:
            self.status.set(f"{acc} — {n_ex} éxons importados.")


# ══════════════════════════════════════════════════════════════
# MODO MANUAL — Upload .docx
# ══════════════════════════════════════════════════════════════

class ModoManual(tk.Toplevel, EditorMixin):
    def __init__(self, parent):
        super().__init__(parent)
        self.title('🧬 ExonEditor v7.3 — Modo Manual (Upload .docx)')
        self.geometry('960x840'); self.minsize(840,640)
        self.configure(bg=C['fundo'])

        self.sequencia_verificada = None
        self.exons                = []
        self.cor_base             = '#aaaaaa'
        self.cor_exon_atual       = '#000000'
        self.fonte_base           = tk.StringVar(value='Courier New')
        self.tamanho_base         = tk.IntVar(value=11)
        self.fonte_exon           = tk.StringVar(value='Courier New')
        self.tamanho_exon         = tk.IntVar(value=11)
        self.caixa                = tk.StringVar(value='minuscula')
        self.arquivo_entrada      = tk.StringVar()
        self._codon_grifo         = None
        self.cor_grifo            = '#7c3aed'
        self._utr_regiao          = None
        self.cor_utr              = '#6699cc'
        self._stop_grifo          = None
        self.cor_stop             = '#cc0000'
        self.marc_inicio          = None
        self.marc_stop            = None
        self.marc_loc             = None
        self._build()

    def _janela_ref(self): return self

    def _build(self):
        hdr = tk.Frame(self, bg='#1a3a1a'); hdr.pack(fill=tk.X)
        tk.Label(hdr, text='📄  Modo Manual — Upload .docx',
                 font=('Arial',14,'bold'), bg='#1a3a1a', fg='white',
                 pady=10, padx=16).pack(side=tk.LEFT)
        tk.Button(hdr, text='← Voltar ao Início', command=self._voltar,
                  bg='#2d5a2d', fg='white', font=('Arial',9),
                  relief=tk.FLAT, padx=10, pady=6,
                  cursor='hand2').pack(side=tk.RIGHT, padx=12, pady=8)

        outer = tk.Frame(self, bg=C['fundo']); outer.pack(fill=tk.BOTH, expand=True)
        cv = tk.Canvas(outer, bg=C['fundo'], highlightthickness=0)
        vsb = ttk.Scrollbar(outer, orient=tk.VERTICAL, command=cv.yview)
        cv.configure(yscrollcommand=vsb.set)
        vsb.pack(side=tk.RIGHT, fill=tk.Y); cv.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.inner = tk.Frame(cv, bg=C['fundo'])
        wid = cv.create_window((0,0), window=self.inner, anchor='nw')
        cv.bind('<Configure>', lambda e: cv.itemconfig(wid, width=e.width))
        self.inner.bind('<Configure>', lambda e: cv.configure(scrollregion=cv.bbox('all')))
        def _scroll_manual(e):
            try: cv.yview_scroll(int(-1*(e.delta/120)),'units')
            except Exception: pass
        cv.bind('<MouseWheel>', _scroll_manual)
        self.inner.bind_all('<MouseWheel>', _scroll_manual)
        # Desvincula ao fechar
        self.protocol('WM_DELETE_WINDOW', lambda: (self.inner.unbind_all('<MouseWheel>'), self._voltar()))

        p = tk.Frame(self.inner, bg=C['fundo'], padx=18, pady=12)
        p.pack(fill=tk.BOTH, expand=True)

        self._build_arquivo(p)
        self._build_checkpoint(p)
        self._build_fmt_base(p)
        self._build_tabela_exons(p)
        self._build_barra(p)
        self._build_codons(p)
        self._build_btns_gerar(p)

    def _build_arquivo(self, p):
        c = self._card(p, '① Arquivo de Entrada  (.docx com a sequência genética)',
                       cor_t='white', bg='#1a3a1a', pady_top=0)
        c.master.configure(highlightbackground='#2d5a2d')
        fr = tk.Frame(c, bg='#1a3a1a'); fr.pack(fill=tk.X)
        self.entry_arq = tk.Entry(fr, textvariable=self.arquivo_entrada,
                                   font=('Arial',10), width=52,
                                   relief=tk.SOLID, bd=1, bg='#1e3a1e', fg='white',
                                   insertbackground='white')
        self.entry_arq.pack(side=tk.LEFT, padx=(0,8), ipady=4)
        tk.Button(fr, text='📂  Selecionar .docx', command=self._selecionar,
                  bg='#15803d', fg='white', font=('Arial',9,'bold'),
                  relief=tk.FLAT, padx=12, pady=5, cursor='hand2').pack(side=tk.LEFT)

    def _build_checkpoint(self, p):
        fr_btn = tk.Frame(p, bg=C['fundo']); fr_btn.pack(fill=tk.X, pady=(10,0))
        self._btn(fr_btn, '🔍  Verificar Sequência  →  Checkpoint',
                  self._verificar, C['amber'], size=10, padx=16, pady=8, anchor=None)

        self.fr_cp_wrap = tk.Frame(p, bg=C['fundo']); self.fr_cp_wrap.pack(fill=tk.X)
        self.fr_cp_wrap.columnconfigure(0, weight=1)
        self.frame_cp = tk.Frame(self.fr_cp_wrap, bg=C['amber_bg'],
                                  highlightbackground=C['amber'], highlightthickness=1,
                                  padx=14, pady=10)
        self.frame_cp.grid(row=0, column=0, sticky='ew', pady=(6,0))
        self.frame_cp.grid_remove()
        self.lbl_cp_total = tk.Label(self.frame_cp, text='', font=('Arial',12,'bold'),
                                      bg=C['amber_bg'], fg=C['amber'])
        self.lbl_cp_total.pack(anchor=tk.W)
        self.lbl_cp_info = tk.Label(self.frame_cp, text='', font=('Arial',8),
                                     bg=C['amber_bg'], fg=C['sub'], justify=tk.LEFT)
        self.lbl_cp_info.pack(anchor=tk.W, pady=(2,6))
        tk.Label(self.frame_cp, text='Prévia (120 nt):',
                 font=('Arial',8,'bold'), bg=C['amber_bg'], fg='#555').pack(anchor=tk.W)
        self.txt_prev = tk.Text(self.frame_cp, height=2, font=('Courier New',9),
                                bg='#fffde7', state=tk.DISABLED, relief=tk.FLAT, bd=1)
        self.txt_prev.pack(fill=tk.X, pady=(2,8))
        fr_cp_btn = tk.Frame(self.frame_cp, bg=C['amber_bg']); fr_cp_btn.pack(anchor=tk.W)
        self._btn(fr_cp_btn,'✅  Confirmar e Prosseguir',self._confirmar,C['green'],size=10,padx=14,pady=6)
        self._btn(fr_cp_btn,'↩  Resetar',self._resetar,C['cinza'],size=9,padx=10,pady=6)

    def _selecionar(self):
        path = filedialog.askopenfilename(filetypes=[('Word','*.docx')])
        if path:
            self.arquivo_entrada.set(path)
            self._resetar(silencioso=True)
            self.status.set(f'Arquivo: {os.path.basename(path)}')

    def _verificar(self):
        if not self.arquivo_entrada.get():
            messagebox.showerror('Erro','Selecione um arquivo .docx.'); return
        self.status.set('Lendo sequência…'); self.update()
        try:
            doc = Document(self.arquivo_entrada.get())
            seq = limpar_sequencia(doc)
        except Exception as e:
            messagebox.showerror('Erro',str(e)); return
        if not seq: messagebox.showerror('Erro','Nenhuma sequência ATCG encontrada.'); return
        total = len(seq)
        self.sequencia_verificada = seq
        self.lbl_cp_total.config(text=f'📊  Total na sequência limpa:  {total:,} pb')
        self.lbl_cp_info.config(
            text=f'Arquivo: {os.path.basename(self.arquivo_entrada.get())}\n'
                 f'Números, espaços e marcas de parágrafo removidos.')
        self.txt_prev.config(state=tk.NORMAL, bg='#fffde7')
        self.txt_prev.delete('1.0',tk.END); self.txt_prev.insert(tk.END, seq[:120])
        self.txt_prev.config(state=tk.DISABLED)
        self.frame_cp.config(bg=C['amber_bg'])
        self.lbl_cp_total.config(bg=C['amber_bg'], fg=C['amber'])
        self.lbl_cp_info.config(bg=C['amber_bg'])
        self.frame_cp.grid()
        self.status.set(f'Checkpoint: {total:,} pb — confirme para prosseguir.')

    def _confirmar(self):
        if not self.sequencia_verificada: return
        total = len(self.sequencia_verificada)
        self.frame_cp.config(bg=C['green_bg'])
        self.lbl_cp_total.config(bg=C['green_bg'], fg=C['green'],
                                  text=f'✅  Sequência confirmada:  {total:,} pb')
        self.lbl_cp_info.config(bg=C['green_bg'], text='Pronto para adicionar éxons.')
        self.txt_prev.config(bg='#dcfce7')
        self._habilitar_exons()
        self.status.set('Sequência confirmada.')

    def _resetar(self, silencioso=False):
        self.sequencia_verificada = None
        self.frame_cp.grid_remove()
        self._desabilitar_exons()
        if not silencioso: self.status.set('Resetado.')


# ══════════════════════════════════════════════════════════════
# TELA INICIAL
# ══════════════════════════════════════════════════════════════

class TelaInicial:
    def __init__(self, root):
        self.root = root
        root.title('🧬 ExonEditor v7.3')
        root.geometry('700x520')
        root.resizable(False, False)
        root.configure(bg='#0d1b2e')
        self._build()

    def _build(self):
        BG = '#0d1b2e'

        # ── Logotipo: Alfa-Hélice ──
        fr_top = tk.Frame(self.root, bg=BG)
        fr_top.pack(pady=(32, 0))

        logo = tk.Canvas(fr_top, width=100, height=90, bg=BG, highlightthickness=0)
        logo.pack()
        self._draw_alfa_helice(logo)

        tk.Label(fr_top, text='ExonEditor', font=('Arial',28,'bold'),
                 bg=BG, fg='white').pack(pady=(6,0))
        tk.Label(fr_top, text='Anotação e análise de Éxons em sequências genéticas',
                 font=('Arial',11), bg=BG, fg='#94a3b8').pack(pady=(4,0))

        tk.Frame(self.root, bg='#1e3a5c', height=1).pack(fill=tk.X, padx=40, pady=22)

        # ── Cartões ──
        fr_cards = tk.Frame(self.root, bg=BG)
        fr_cards.pack(padx=40)

        self._card_modo(fr_cards, tk.LEFT,
            titulo='Modo Automático', subtitulo='Integração NCBI API',
            descricao='Busca, sequência e éxons\nimportados automaticamente\ndo GenBank sem copiar nada.',
            icon_fn=self._icon_api,
            icon_bg='#1e4a7a', card_bg='#0f2744', borda='#1e4a7a',
            btn_cmd=self._abrir_automatico)

        tk.Frame(fr_cards, bg='#1e3a5c', width=1).pack(side=tk.LEFT, fill=tk.Y, padx=14)

        self._card_modo(fr_cards, tk.LEFT,
            titulo='Modo Manual', subtitulo='Upload de arquivo .docx',
            descricao='Faça upload do seu arquivo\nWord com a sequência e cole\nos dados do GenBank.',
            icon_fn=self._icon_docx,
            icon_bg='#1e3a5c', card_bg='#0d1b38', borda='#2d4a8a',
            btn_cmd=self._abrir_manual)

        tk.Label(self.root,
                 text='Ambos os modos incluem: Barra visual  •  Análise de códons  •  Proteínas  •  Geração de documentos Word',
                 font=('Arial',8), bg=BG, fg='#475569').pack(pady=(18,0))

    def _draw_alfa_helice(self, cv):
        """Logotipo: Alfa-hélice estilizada — estrutura em espiral com ribbon."""
        W, H = 100, 90
        import math
        # Ribbon da hélice: série de elipses deslocadas simulando espiral 3D
        cores = ['#4ade80','#34d399','#2563eb','#60a5fa','#4ade80','#34d399','#2563eb']
        for i in range(7):
            t = i / 6.0
            cx = 50 + 22 * math.sin(t * math.pi * 2.2)
            cy = 10 + t * 70
            w  = 28 - abs(math.sin(t * math.pi * 2.2)) * 8
            h  = 14
            cor = cores[i]
            # Sombra
            cv.create_oval(cx-w//2+2, cy-h//2+2, cx+w//2+2, cy+h//2+2,
                           fill='#0a1628', outline='')
            # Ribbon
            cv.create_oval(cx-w//2, cy-h//2, cx+w//2, cy+h//2,
                           fill=cor, outline='white', width=1)
        # Eixo central
        cv.create_line(50, 8, 50, 82, fill='#475569', width=1, dash=(3,3))

    def _card_modo(self, parent, lado, titulo, subtitulo, descricao,
                   icon_fn, icon_bg, card_bg, borda, btn_cmd):
        fr = tk.Frame(parent, bg=card_bg,
                      highlightbackground=borda, highlightthickness=1,
                      padx=22, pady=18)
        fr.pack(side=lado, fill=tk.BOTH, expand=True)

        tk.Label(fr, text=titulo, font=('Arial',13,'bold'),
                 bg=card_bg, fg='white').pack()
        tk.Label(fr, text=subtitulo, font=('Arial',9,'italic'),
                 bg=card_bg, fg='#94a3b8').pack(pady=(2,6))
        tk.Label(fr, text=descricao, font=('Arial',9),
                 bg=card_bg, fg='#cbd5e1', justify=tk.CENTER).pack(pady=(0,14))

        # Ícone GRANDE como botão principal
        cv = tk.Canvas(fr, width=130, height=130, bg=icon_bg,
                       highlightthickness=3, highlightbackground=borda,
                       cursor='hand2')
        cv.pack()
        icon_fn(cv, icon_bg)

        def _on_enter(e):
            cv.configure(highlightbackground='white', highlightthickness=4)
        def _on_leave(e):
            cv.configure(highlightbackground=borda, highlightthickness=3)
        def _on_click(e): btn_cmd()

        cv.bind('<Enter>', _on_enter)
        cv.bind('<Leave>', _on_leave)
        cv.bind('<Button-1>', _on_click)

        tk.Label(fr, text='Clique para iniciar', font=('Arial',8,'italic'),
                 bg=card_bg, fg='#64748b', cursor='hand2').pack(pady=(6,0))

    def _icon_api(self, cv, bg):
        """Globo/API — 130x130, azul escuro"""
        cv.create_oval(10,10,120,120, fill='#1e3a5c', outline='#60a5fa', width=2)
        cv.create_line(65,10,65,120, fill='#93c5fd', width=1)
        cv.create_line(10,65,120,65, fill='#93c5fd', width=1)
        # Meridianos curvos simulados
        for dx in [-24, 0, 24]:
            cv.create_oval(65+dx-18,10, 65+dx+18,120, outline='#3b82f6', width=1, fill='')
        cv.create_oval(24,24,106,106, outline='#bfdbfe', width=1, fill='')
        cv.create_oval(44,44,86,86,  fill='#2563eb', outline='#93c5fd', width=2)
        cv.create_text(65,65, text='⚡', font=('Arial',20,'bold'), fill='white')
        cv.create_text(65,112, text='NCBI API', font=('Arial',8,'bold'), fill='#93c5fd')

    def _icon_docx(self, cv, bg):
        """Documento .docx — 130x130, azul médio"""
        # Papel
        cv.create_rectangle(22,8,108,118, fill='#1e3a7a', outline='#60a5fa', width=2)
        # Cabeçalho
        cv.create_rectangle(22,8,108,34, fill='#2563eb', outline='')
        cv.create_text(65,21, text='.docx', font=('Arial',12,'bold'), fill='white')
        # Linhas de texto
        for y in [46,58,70,82,94]:
            w = 68 if y != 82 else 44
            cv.create_rectangle(30,y, 30+w,y+7, fill='#60a5fa', outline='')
        # Seta de upload
        cv.create_oval(82,88,108,114, fill='#2563eb', outline='#93c5fd', width=2)
        cv.create_text(95,101, text='↑', font=('Arial',14,'bold'), fill='white')

    def _abrir_automatico(self):
        self.root.withdraw()
        win = ModoAutomatico(self.root)

    def _abrir_manual(self):
        self.root.withdraw()
        win = ModoManual(self.root)


# ══════════════════════════════════════════════════════════════
if __name__ == '__main__':
    root = tk.Tk()
    TelaInicial(root)
    root.mainloop()

"""
ExonEditor Web — Backend (FastAPI)
Processa arquivos .docx com sequências genéticas e retorna documento anotado.
"""

from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
import re, os, uuid, tempfile, json
from docx import Document
from docx.shared import RGBColor, Pt

app = FastAPI(title="ExonEditor API")

# Permite requisições do frontend (qualquer origem em dev; restrinja em produção)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# ──────────────────────────────────────────────
# FUNÇÕES DE PROCESSAMENTO (idênticas ao desktop)
# ──────────────────────────────────────────────

def limpar_e_extrair_sequencia(doc):
    texto = ""
    for para in doc.paragraphs:
        t = re.sub(r'\d+', '', para.text)
        t = re.sub(r'[\s\r\n\v\f\u000b]+', '', t)
        texto += t
    return re.sub(r'[^ATCGatcg]', '', texto).upper()


def parsear_ncbi(texto):
    exons = []
    padrao     = re.compile(r'exon\s+(\d+)\.\.(\d+)', re.IGNORECASE)
    padrao_num = re.compile(r'/number=(\d+)')
    matches_pos = list(padrao.finditer(texto))
    matches_num = list(padrao_num.finditer(texto))
    for i, m in enumerate(matches_pos):
        numero = int(matches_num[i].group(1)) if i < len(matches_num) else i + 1
        exons.append({'numero': numero, 'inicio': int(m.group(1)), 'fim': int(m.group(2))})
    return sorted(exons, key=lambda x: x['numero'])


def cor_hex_para_rgb(hex_color):
    h = hex_color.lstrip('#')
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))


def gerar_documento_anotado(sequencia, config_base, exons, output_path,
                             chars_por_linha=60, caixa='maiuscula'):
    sequencia = sequencia.lower() if caixa == 'minuscula' else sequencia.upper()

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin    = Pt(36)
    sec.bottom_margin = Pt(36)
    sec.left_margin   = Pt(54)
    sec.right_margin  = Pt(54)

    r_b, g_b, b_b = cor_hex_para_rgb(config_base['cor'])
    n = len(sequencia)

    anotacoes = [None] * n
    for exon in exons:
        ini = max(0, exon['inicio'] - 1)
        fim = min(n, exon['fim'])
        for i in range(ini, fim):
            anotacoes[i] = exon

    for linha_idx in range(0, n, chars_por_linha):
        trecho   = sequencia[linha_idx : linha_idx + chars_por_linha]
        anot_lin = anotacoes[linha_idx : linha_idx + chars_por_linha]
        para     = doc.add_paragraph()
        para.paragraph_format.space_after  = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        para.style.font.name = config_base['fonte']

        i = 0
        while i < len(trecho):
            cur = anot_lin[i]
            j = i + 1
            while j < len(trecho) and anot_lin[j] is cur:
                j += 1
            run = para.add_run(trecho[i:j])
            if cur is None:
                run.bold           = False
                run.font.name      = config_base['fonte']
                run.font.size      = Pt(config_base['tamanho'])
                run.font.color.rgb = RGBColor(r_b, g_b, b_b)
            else:
                r_e, g_e, b_e     = cor_hex_para_rgb(cur['cor'])
                run.bold           = True
                run.font.name      = cur['fonte']
                run.font.size      = Pt(cur['tamanho'])
                run.font.color.rgb = RGBColor(r_e, g_e, b_e)
            i = j

    doc.save(output_path)
    return n


# ──────────────────────────────────────────────
# ENDPOINTS
# ──────────────────────────────────────────────

@app.get("/")
def raiz():
    return {"status": "ExonEditor API online"}


@app.post("/verificar")
async def verificar_sequencia(file: UploadFile = File(...)):
    """
    Recebe o .docx, limpa e retorna:
    - total de nucleotídeos
    - prévia dos primeiros 120 nt
    - sequência completa (para uso no frontend sem re-upload)
    """
    if not file.filename.endswith('.docx'):
        raise HTTPException(400, "Envie um arquivo .docx válido.")

    conteudo = await file.read()
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    tmp.write(conteudo)
    tmp.close()

    try:
        doc = Document(tmp.name)
        sequencia = limpar_e_extrair_sequencia(doc)
    finally:
        os.unlink(tmp.name)

    if not sequencia:
        raise HTTPException(422, "Nenhuma sequência ATCG encontrada no documento.")

    return {
        "total": len(sequencia),
        "previa": sequencia[:120],
        "sequencia": sequencia   # retorna para o frontend armazenar em estado
    }


@app.post("/parsear-ncbi")
async def parsear_ncbi_endpoint(texto: str = Form(...)):
    """Parseia o texto colado do NCBI e retorna lista de Exons."""
    exons = parsear_ncbi(texto)
    if not exons:
        raise HTTPException(422, "Nenhum Exon encontrado. Verifique o formato do texto NCBI.")
    return {"exons": exons, "total": len(exons)}


@app.post("/gerar")
async def gerar_documento(
    sequencia:      str  = Form(...),
    config_base:    str  = Form(...),   # JSON string
    exons:          str  = Form(...),   # JSON string
    caixa:          str  = Form("maiuscula"),
    chars_por_linha: int = Form(60),
):
    """
    Recebe a sequência limpa + configurações e retorna o .docx anotado para download.
    """
    try:
        cfg  = json.loads(config_base)
        exns = json.loads(exons)
    except Exception:
        raise HTTPException(400, "Parâmetros inválidos.")

    if not sequencia:
        raise HTTPException(400, "Sequência vazia.")

    # Reconstrói objetos de exon com referências únicas (necessário para o agrupamento)
    exns_obj = []
    for e in exns:
        exns_obj.append({
            'inicio':  int(e['inicio']),
            'fim':     int(e['fim']),
            'fonte':   e.get('fonte', 'Courier New'),
            'tamanho': int(e.get('tamanho', 14)),
            'cor':     e.get('cor', '#000000'),
            'numero':  e.get('numero', 0),
        })

    tmp_out = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    tmp_out.close()

    try:
        total = gerar_documento_anotado(
            sequencia, cfg, exns_obj, tmp_out.name,
            chars_por_linha=chars_por_linha, caixa=caixa
        )
    except Exception as e:
        os.unlink(tmp_out.name)
        raise HTTPException(500, f"Erro ao gerar documento: {str(e)}")

    nome_arquivo = f"sequencia_anotada_{uuid.uuid4().hex[:6]}.docx"

    return FileResponse(
        tmp_out.name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=nome_arquivo,
        background=None
    )
